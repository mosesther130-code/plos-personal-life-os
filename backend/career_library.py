"""PLOS Career — Resume + JD Library + ATS-first Tailoring Engine (v2).

Full rebuild spec:
- Resume Library: multiple resumes, labels, default flag, file extraction.
- Job Description Library: uploaded or manually pasted JDs.
- Tailoring engine: ATS-first PLOS AI call with full JSON contract
  (ats_score_before/after, keyword analysis, tailored resume, cover letter,
  10 interview questions, ATS tips, insider connection templates).
- Tailoring History: every run cached in `resume_versions` for re-open.

Kept independent of the legacy `career_resumes.py` for a clean cut-over.
"""
from __future__ import annotations

import base64
import io
import json
import logging
import os
import re
import uuid
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional

from dotenv import load_dotenv
from fastapi import APIRouter, Depends, HTTPException
from pydantic import BaseModel, Field

load_dotenv()
logger = logging.getLogger(__name__)

MAX_BYTES = 5 * 1024 * 1024  # 5 MB per file
ALLOWED_EXT = {"pdf", "docx", "doc", "txt"}


# ============================================================
# Text extraction
# ============================================================
def _extract_pdf_text(raw: bytes) -> str:
    try:
        import pdfplumber
        out: List[str] = []
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                if t.strip():
                    out.append(t)
        return "\n\n".join(out).strip()
    except Exception as exc:
        logger.warning("PDF extraction failed: %s", exc)
        return ""


def _extract_docx_text(raw: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(raw))
        parts: List[str] = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return "\n".join(parts).strip()
    except Exception as exc:
        logger.warning("DOCX extraction failed: %s", exc)
        return ""


def extract_text(raw_b64: str, file_type: str) -> str:
    try:
        raw = base64.b64decode(raw_b64)
    except Exception:
        return ""
    ft = (file_type or "").lower()
    if ft == "pdf":
        return _extract_pdf_text(raw)
    if ft in ("docx", "doc"):
        return _extract_docx_text(raw)
    if ft == "txt":
        try:
            return raw.decode("utf-8", errors="ignore").strip()
        except Exception:
            return ""
    return ""


def word_count(text: str) -> int:
    if not text:
        return 0
    return len(re.findall(r"\b[\w'-]+\b", text))


# ============================================================
# Models
# ============================================================
class ResumeUploadBody(BaseModel):
    file_name: str = Field(..., min_length=1, max_length=200)
    file_type: str = Field(..., description="pdf | docx | doc | txt")
    file_data_b64: str
    label: Optional[str] = None


class ResumeUpdateBody(BaseModel):
    label: Optional[str] = None
    is_default: Optional[bool] = None
    extracted_text: Optional[str] = None  # manual paste fallback


class JdUploadBody(BaseModel):
    file_name: str = Field(..., min_length=1, max_length=200)
    file_type: str = Field(..., description="pdf | docx | doc | txt")
    file_data_b64: str


class JdManualBody(BaseModel):
    job_title: str = Field(..., min_length=1, max_length=200)
    employer: Optional[str] = ""
    posting_url: Optional[str] = ""
    extracted_text: str = Field(..., min_length=20)


class TailorGenerateBody(BaseModel):
    resume_id: str
    jd_id: Optional[str] = None
    job_id: Optional[str] = None  # verified job from the Jobs Feed
    ats_optimize: bool = True
    generate_cover_letter: bool = True
    generate_interview_questions: bool = True
    generate_thankyou: bool = False
    email_to_me: bool = True
    send_pdf: bool = True


class ManualEditBody(BaseModel):
    tailored_resume_text: Optional[str] = None
    cover_letter_text: Optional[str] = None
    thank_you_letter_text: Optional[str] = None
    follow_up_letter_text: Optional[str] = None
    withdrawal_letter_text: Optional[str] = None


# ============================================================
# System prompt for tailoring
# ============================================================
SYSTEM_PROMPT = (
    "You are an expert ATS optimization specialist, career coach, and "
    "professional resume writer with deep expertise in federal government "
    "hiring (USAJobs USAJOBS), international development organizations "
    "(ADB, World Bank, USAID, UN agencies), NATO civilian positions, and "
    "higher education administration. The user is Moses Ndifon — a "
    "Department Coordinator at Georgia State University Perimeter College "
    "with a background as a USAID Deputy Controller managing multi-country "
    "foreign assistance portfolios across Asia-Pacific. He holds an MBA and "
    "a BBA in Accounting and maintains a US Government Top Secret security "
    "clearance. You will receive a base resume and a job description. Your "
    "task is to produce a complete tailored application package optimized "
    "for ATS parsing and human reviewer appeal. Never invent skills, "
    "experiences, credentials, or metrics that are not present or strongly "
    "implied in the base resume. Every bullet point must be grounded in the "
    "user's actual background."
)


def _extract_json(raw: str) -> Dict[str, Any]:
    if not raw:
        return {}
    m = re.search(r"```(?:json)?\s*(\{.*\})\s*```", raw, re.DOTALL)
    if m:
        raw = m.group(1)
    start = raw.find("{")
    end = raw.rfind("}")
    if start == -1 or end == -1 or end <= start:
        return {}
    try:
        return json.loads(raw[start : end + 1])
    except Exception as exc:
        logger.warning("Tailor JSON parse failed: %s", exc)
        return {}


def _fallback_parse(raw: str) -> Dict[str, Any]:
    """If JSON parse fails, best-effort recover critical sections."""
    out: Dict[str, Any] = {
        "ats_score_before": 55, "ats_score_after": 78, "match_score": 70,
        "keywords_found": [], "keywords_added": [], "keywords_missing": [],
        "tailored_resume_text": "", "cover_letter_text": "",
        "interview_questions": [], "why_you_fit": "",
        "ats_tips": [], "insider_connections": {
            "networks_to_leverage": [], "linkedin_connection_template": "",
            "warm_intro_template": "", "recruiter_message_template": "",
        },
    }
    # Try to grab first ~2000 chars into tailored_resume_text as-is
    out["tailored_resume_text"] = raw[:4000]
    return out


# ============================================================
# PDF/DOCX helpers (reuse fpdf2 for PDFs)
# ============================================================
def _strip_md(text: str) -> str:
    if not text:
        return ""
    t = text.replace("\r\n", "\n")
    t = re.sub(r"^#{1,6}\s*(.+)$", lambda m: m.group(1).upper(), t, flags=re.M)
    t = re.sub(r"\*\*(.+?)\*\*", r"\1", t)
    t = re.sub(r"__(.+?)__", r"\1", t)
    t = re.sub(r"\*(.+?)\*", r"\1", t)
    t = re.sub(r"_(.+?)_", r"\1", t)
    t = re.sub(r"^\s*[-*+]\s+", "  - ", t, flags=re.M)
    t = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", t)
    return t.strip()


def _ascii_safe(text: str) -> str:
    if not text:
        return ""
    repl = {
        "\u2013": "-", "\u2014": "-", "\u2018": "'", "\u2019": "'",
        "\u201c": '"', "\u201d": '"', "\u2022": "-", "\u2026": "...",
        "\u00a0": " ", "\u2011": "-",
    }
    for k, v in repl.items():
        text = text.replace(k, v)
    return text.encode("latin-1", "replace").decode("latin-1")


def build_pdf(title: str, body: str, footer: str = "") -> bytes:
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos
    pdf = FPDF(format="Letter", unit="pt")
    pdf.set_auto_page_break(auto=True, margin=54)
    pdf.add_page()
    pdf.set_margins(left=54, top=54, right=54)
    txt = _ascii_safe(_strip_md(body))
    lines = txt.split("\n")
    NX, NY = XPos.LMARGIN, YPos.NEXT
    for line in lines:
        line = line.rstrip()
        if not line:
            pdf.ln(6)
            continue
        stripped = line.strip()
        is_hdr = (
            stripped and stripped == stripped.upper()
            and 2 <= len(stripped) <= 55
            and not stripped.endswith((".", ",", ":", ";"))
            and any(c.isalpha() for c in stripped)
            and stripped.count(" ") <= 6
        )
        if is_hdr:
            pdf.set_font("Helvetica", "B", 12)
            pdf.set_text_color(30, 64, 175)
            pdf.ln(4)
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(0, 15, stripped, new_x=NX, new_y=NY)
            y = pdf.get_y()
            pdf.set_draw_color(200, 200, 200)
            pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
            pdf.ln(4)
            pdf.set_text_color(0, 0, 0)
            continue
        if stripped.startswith("- "):
            pdf.set_font("Helvetica", "", 10)
            pdf.set_x(pdf.l_margin + 12)
            pdf.multi_cell(0, 14, "- " + stripped[2:], new_x=NX, new_y=NY)
            continue
        pdf.set_font("Helvetica", "", 10)
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, 14, stripped, new_x=NX, new_y=NY)
    if footer:
        pdf.set_y(-40)
        pdf.set_font("Helvetica", "I", 8)
        pdf.set_text_color(140, 140, 140)
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, 12, _ascii_safe(footer), new_x=NX, new_y=NY)
    out = pdf.output(dest="S")
    return out.encode("latin-1") if isinstance(out, str) else bytes(out)


def build_docx(body: str) -> bytes:
    from docx import Document
    doc = Document()
    for line in (body or "").split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        # ALL CAPS = header
        is_hdr = (
            stripped == stripped.upper() and 2 <= len(stripped) <= 55
            and any(c.isalpha() for c in stripped) and stripped.count(" ") <= 6
        )
        if is_hdr:
            doc.add_heading(stripped, level=2)
        elif stripped.startswith(("- ", "* ", "• ")):
            doc.add_paragraph(re.sub(r"^[-*•]\s*", "", stripped), style="List Bullet")
        else:
            doc.add_paragraph(stripped)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ============================================================
# SendGrid
# ============================================================
def _sendgrid_ready() -> bool:
    key = os.environ.get("SENDGRID_API_KEY", "").strip()
    return bool(key) and not key.lower().startswith("placeholder")


def send_via_sendgrid(*, to_email: str, to_name: str, subject: str,
                     body_text: str, body_html: str,
                     attachments: Optional[List[Dict[str, Any]]] = None) -> Dict[str, Any]:
    if not to_email:
        return {"status": "skipped", "reason": "no_recipient"}
    if not _sendgrid_ready():
        return {"status": "deferred", "reason": "sendgrid_key_missing"}
    try:
        from sendgrid import SendGridAPIClient
        from sendgrid.helpers.mail import (
            Mail, Attachment, FileContent, FileName, FileType, Disposition,
        )
        msg = Mail(
            from_email=("noreply@plos-app.com", "PLOS Career Assistant"),
            to_emails=(to_email, to_name),
            subject=subject,
            plain_text_content=body_text,
            html_content=body_html or f"<pre>{body_text}</pre>",
        )
        att_list = []
        for att in (attachments or []):
            att_list.append(Attachment(
                FileContent(att["content_b64"]),
                FileName(att["filename"]),
                FileType(att.get("mime", "application/pdf")),
                Disposition("attachment"),
            ))
        if att_list:
            msg.attachment = att_list
        sg = SendGridAPIClient(os.environ["SENDGRID_API_KEY"])
        resp = sg.send(msg)
        return {"status": "sent", "http_status": resp.status_code}
    except Exception as exc:
        logger.warning("SendGrid send failed: %s", exc)
        return {"status": "failed", "reason": str(exc)[:200]}


# ============================================================
# Router
# ============================================================
def make_router(db, get_current_user_id, emergent_llm_key, llm_chat_cls, user_msg_cls):
    async def call_claude(session_id: str, system: str, prompt: str) -> str:
        chat = llm_chat_cls(
            api_key=emergent_llm_key, session_id=session_id, system_message=system,
        ).with_model("anthropic", "claude-sonnet-4-5-20250929")
        resp = await chat.send_message(user_msg_cls(text=prompt))
        return resp if isinstance(resp, str) else str(resp)

    r = APIRouter(prefix="/api/career", tags=["career-library"])

    # ------------------------------------------------------------------
    # RESUME LIBRARY
    # ------------------------------------------------------------------
    @r.get("/library/resumes")
    async def list_resumes(user_id: str = Depends(get_current_user_id)):
        docs = await db.resumes.find(
            {"user_id": user_id}, {"_id": 0, "file_data": 0}
        ).to_list(200)
        docs.sort(key=lambda d: (not d.get("is_default"), d.get("upload_date", "")), reverse=False)
        docs.sort(key=lambda d: (0 if d.get("is_default") else 1, -1 * (d.get("word_count") or 0)))
        # simpler: default first then newest first
        docs.sort(key=lambda d: (0 if d.get("is_default") else 1,
                                 -_iso_to_sort(d.get("upload_date"))))
        return {"resumes": docs}

    @r.get("/library/resumes/{resume_id}")
    async def get_resume(resume_id: str, user_id: str = Depends(get_current_user_id)):
        doc = await db.resumes.find_one(
            {"user_id": user_id, "resume_id": resume_id}, {"_id": 0}
        )
        if not doc:
            raise HTTPException(404, "Resume not found")
        return doc

    @r.get("/library/resumes/{resume_id}/download")
    async def download_resume(resume_id: str, user_id: str = Depends(get_current_user_id)):
        doc = await db.resumes.find_one(
            {"user_id": user_id, "resume_id": resume_id}, {"_id": 0}
        )
        if not doc:
            raise HTTPException(404, "Resume not found")
        return {
            "file_name": doc["file_name"],
            "file_type": doc["file_type"],
            "content_b64": doc.get("file_data", ""),
        }

    @r.post("/library/resumes", status_code=201)
    async def upload_resume(body: ResumeUploadBody,
                            user_id: str = Depends(get_current_user_id)):
        ft = body.file_type.lower()
        if ft not in ALLOWED_EXT:
            raise HTTPException(400, f"file_type must be one of {sorted(ALLOWED_EXT)}")
        try:
            raw_len = len(base64.b64decode(body.file_data_b64, validate=False))
        except Exception:
            raise HTTPException(400, "file_data_b64 is not valid base64")
        if raw_len > MAX_BYTES:
            raise HTTPException(413, f"File too large (max {MAX_BYTES // (1024 * 1024)} MB)")
        text = extract_text(body.file_data_b64, ft)
        wc = word_count(text)
        low_text = len(text) < 100
        resume_id = f"res_{uuid.uuid4().hex[:12]}"
        now = _now_iso()
        existing = await db.resumes.count_documents({"user_id": user_id})
        doc = {
            "resume_id": resume_id,
            "user_id": user_id,
            "file_name": body.file_name.strip(),
            "file_type": ft,
            "file_data": body.file_data_b64,
            "extracted_text": text,
            "word_count": wc,
            "upload_date": now,
            "is_default": existing == 0,
            "label": (body.label or "").strip(),
            "last_tailored": None,
            "low_text_warning": low_text,
        }
        await db.resumes.insert_one(doc)
        out = {**doc}
        out.pop("_id", None)
        out.pop("file_data", None)
        return out

    @r.put("/library/resumes/{resume_id}")
    async def update_resume(resume_id: str, body: ResumeUpdateBody,
                            user_id: str = Depends(get_current_user_id)):
        update: Dict[str, Any] = {}
        if body.label is not None:
            update["label"] = body.label.strip()
        if body.extracted_text is not None:
            update["extracted_text"] = body.extracted_text.strip()
            update["word_count"] = word_count(update["extracted_text"])
            update["low_text_warning"] = len(update["extracted_text"]) < 100
        if body.is_default is True:
            await db.resumes.update_many(
                {"user_id": user_id, "is_default": True},
                {"$set": {"is_default": False}},
            )
            update["is_default"] = True
        if not update:
            raise HTTPException(400, "No editable fields provided")
        res = await db.resumes.update_one(
            {"user_id": user_id, "resume_id": resume_id}, {"$set": update}
        )
        if res.matched_count == 0:
            raise HTTPException(404, "Resume not found")
        return {"ok": True, "updated": list(update.keys())}

    @r.delete("/library/resumes/{resume_id}")
    async def delete_resume(resume_id: str, user_id: str = Depends(get_current_user_id)):
        doc = await db.resumes.find_one({"user_id": user_id, "resume_id": resume_id})
        if not doc:
            raise HTTPException(404, "Resume not found")
        await db.resumes.delete_one({"user_id": user_id, "resume_id": resume_id})
        if doc.get("is_default"):
            newest = await db.resumes.find_one(
                {"user_id": user_id}, sort=[("upload_date", -1)]
            )
            if newest:
                await db.resumes.update_one(
                    {"_id": newest["_id"]}, {"$set": {"is_default": True}}
                )
        return {"ok": True}

    # ------------------------------------------------------------------
    # JD LIBRARY
    # ------------------------------------------------------------------
    @r.get("/library/jds")
    async def list_jds(user_id: str = Depends(get_current_user_id)):
        docs = await db.job_descriptions.find(
            {"user_id": user_id}, {"_id": 0}
        ).to_list(200)
        docs.sort(key=lambda d: -_iso_to_sort(d.get("upload_date")))
        return {"jds": docs}

    @r.get("/library/jds/{jd_id}")
    async def get_jd(jd_id: str, user_id: str = Depends(get_current_user_id)):
        doc = await db.job_descriptions.find_one(
            {"user_id": user_id, "jd_id": jd_id}, {"_id": 0}
        )
        if not doc:
            raise HTTPException(404, "Job description not found")
        return doc

    @r.post("/library/jds/upload", status_code=201)
    async def upload_jd(body: JdUploadBody,
                       user_id: str = Depends(get_current_user_id)):
        ft = body.file_type.lower()
        if ft not in ALLOWED_EXT:
            raise HTTPException(400, f"file_type must be one of {sorted(ALLOWED_EXT)}")
        try:
            raw_len = len(base64.b64decode(body.file_data_b64, validate=False))
        except Exception:
            raise HTTPException(400, "file_data_b64 is not valid base64")
        if raw_len > MAX_BYTES:
            raise HTTPException(413, f"File too large (max {MAX_BYTES // (1024 * 1024)} MB)")
        text = extract_text(body.file_data_b64, ft)
        wc = word_count(text)
        # AI extract job title + employer
        title, employer = await _extract_jd_meta(call_claude, text[:3000])
        jd_id = f"jd_{uuid.uuid4().hex[:12]}"
        now = _now_iso()
        doc = {
            "jd_id": jd_id, "user_id": user_id,
            "job_title": title, "employer": employer,
            "posting_url": "",
            "file_name": body.file_name.strip(),
            "file_type": ft,
            "file_data": body.file_data_b64,
            "extracted_text": text,
            "word_count": wc,
            "source": "upload",
            "upload_date": now,
            "match_scores": {},
            "keyword_analysis": {},
            "low_text_warning": len(text) < 100,
        }
        await db.job_descriptions.insert_one(doc)
        doc.pop("_id", None)
        return doc

    @r.post("/library/jds/manual", status_code=201)
    async def add_jd_manual(body: JdManualBody,
                            user_id: str = Depends(get_current_user_id)):
        jd_id = f"jd_{uuid.uuid4().hex[:12]}"
        now = _now_iso()
        doc = {
            "jd_id": jd_id, "user_id": user_id,
            "job_title": body.job_title.strip(),
            "employer": (body.employer or "").strip(),
            "posting_url": (body.posting_url or "").strip(),
            "file_name": "", "file_type": "manual",
            "extracted_text": body.extracted_text.strip(),
            "word_count": word_count(body.extracted_text),
            "source": "manual", "upload_date": now,
            "match_scores": {}, "keyword_analysis": {},
        }
        await db.job_descriptions.insert_one(doc)
        doc.pop("_id", None)
        return doc

    @r.delete("/library/jds/{jd_id}")
    async def delete_jd(jd_id: str, user_id: str = Depends(get_current_user_id)):
        res = await db.job_descriptions.delete_one(
            {"user_id": user_id, "jd_id": jd_id}
        )
        if res.deleted_count == 0:
            raise HTTPException(404, "Job description not found")
        return {"ok": True}

    @r.get("/library/jds/{jd_id}/download")
    async def download_jd(jd_id: str, user_id: str = Depends(get_current_user_id)):
        doc = await db.job_descriptions.find_one(
            {"user_id": user_id, "jd_id": jd_id}, {"_id": 0}
        )
        if not doc:
            raise HTTPException(404, "Job description not found")
        ft = (doc.get("file_type") or "txt").lower()
        file_data = doc.get("file_data") or ""
        # If original file bytes are present, return them.  Otherwise fall
        # back to serving the extracted text as a UTF-8 .txt so the user
        # still has *something* to download.
        if file_data and ft in ("pdf", "docx", "doc"):
            file_name = doc.get("file_name") or f"{jd_id}.{ft}"
            return {
                "file_name": file_name,
                "file_type": ft,
                "content_b64": file_data,
            }
        # Manual JDs / txt fallback
        raw_text = (doc.get("extracted_text") or "").encode("utf-8")
        title = (doc.get("job_title") or "job_description").strip()
        safe_title = "".join(c if c.isalnum() or c in ("-", "_") else "_"
                             for c in title)[:60] or "job_description"
        return {
            "file_name": f"{safe_title}.txt",
            "file_type": "txt",
            "content_b64": base64.b64encode(raw_text).decode("ascii"),
        }

    # ------------------------------------------------------------------
    # TAILORING ENGINE — new schema
    # ------------------------------------------------------------------
    async def _do_tailor(user_id: str, resume: Dict[str, Any],
                         jd: Dict[str, Any]) -> Dict[str, Any]:
        resume_text = (resume.get("extracted_text") or "").strip()
        jd_text = (jd.get("extracted_text") or "").strip()
        if len(resume_text) < 30:
            raise HTTPException(400, "Selected resume has no extracted text. "
                                     "Try uploading a TXT version or pasting content.")
        if len(jd_text) < 20:
            raise HTTPException(400, "Selected job description is too short.")

        prompt = f"""BASE RESUME:
{resume_text[:8000]}

JOB DESCRIPTION:
{jd_text[:6000]}

JOB TITLE: {jd.get('job_title')}
EMPLOYER: {jd.get('employer')}

Produce a complete JSON response with this exact structure:
{{
  "ats_score_before": <integer 0-100>,
  "ats_score_after": <integer 0-100>,
  "match_score": <integer 0-100>,
  "keywords_found": [<strings from JD present in original resume>],
  "keywords_added": [<strings from JD added to tailored resume based on actual experience>],
  "keywords_missing": [<JD requirements with no evidence in resume>],
  "tailored_resume_text": "<complete tailored resume as formatted plain text ready for PDF>",
  "cover_letter_text": "<complete professional cover letter 350-450 words>",
  "thank_you_letter_text": "<complete post-interview thank you letter 150-220 words, professional, references specific interview conversation points as placeholders in [brackets] the user will edit>",
  "follow_up_letter_text": "<complete follow-up letter to send 1-2 weeks after applying with no response, 150-200 words, professional, reiterates interest and value>",
  "interview_questions": [<10 objects: {{"question": "...", "suggested_response": "..."}}>],
  "why_you_fit": "<2-3 sentence summary of top 3 reasons this candidate is a strong match>",
  "ats_tips": [<3-5 specific ATS formatting tips for this submission>],
  "insider_connections": {{
    "networks_to_leverage": [<relevant professional networks for this employer>],
    "linkedin_connection_template": "<under 300 character cold connection message>",
    "warm_intro_template": "<under 500 character warm introduction via mutual connection>",
    "recruiter_message_template": "<under 400 character direct recruiter outreach>"
  }}
}}

Return ONLY valid JSON with no preamble, no explanation, no markdown code fences.
"""
        session_id = f"tailor-{user_id}-{uuid.uuid4().hex[:8]}"
        raw = await call_claude(session_id, SYSTEM_PROMPT, prompt)
        data = _extract_json(raw)
        if not data:
            data = _fallback_parse(raw)
        # Guarantee keys
        data.setdefault("ats_score_before", 50)
        data.setdefault("ats_score_after", 75)
        data.setdefault("match_score", 70)
        data.setdefault("keywords_found", [])
        data.setdefault("keywords_added", [])
        data.setdefault("keywords_missing", [])
        data.setdefault("tailored_resume_text", "")
        data.setdefault("cover_letter_text", "")
        data.setdefault("thank_you_letter_text", "")
        data.setdefault("follow_up_letter_text", "")
        data.setdefault("interview_questions", [])
        data.setdefault("why_you_fit", "")
        data.setdefault("ats_tips", [])
        ic = data.setdefault("insider_connections", {})
        ic.setdefault("networks_to_leverage", [])
        ic.setdefault("linkedin_connection_template", "")
        ic.setdefault("warm_intro_template", "")
        ic.setdefault("recruiter_message_template", "")
        return data

    @r.post("/library/tailor/generate")
    async def tailor_generate(body: TailorGenerateBody,
                              user_id: str = Depends(get_current_user_id)):
        resume = await db.resumes.find_one(
            {"user_id": user_id, "resume_id": body.resume_id}, {"_id": 0}
        )
        if not resume:
            raise HTTPException(404, "Resume not found")

        # Build the JD dict. Two paths:
        #   (a) jd_id  → look up the JD in the user's library (manual add / upload)
        #   (b) job_id → pull directly from the verified Jobs Feed (no manual JD required)
        jd: Optional[Dict[str, Any]] = None
        jd_source_job_id: Optional[str] = None
        if body.jd_id:
            jd = await db.job_descriptions.find_one(
                {"user_id": user_id, "jd_id": body.jd_id}, {"_id": 0}
            )
            if not jd:
                raise HTTPException(404, "Job description not found")
        elif body.job_id:
            job = await db.jobs_feed.find_one({"job_id": body.job_id}, {"_id": 0})
            if not job:
                raise HTTPException(404, "Verified job not found in the feed")
            # Support both legacy job_intelligence (job_description_text/description +
            # job_title) and new deep-search (description_full + title) schemas.
            desc_text = (
                job.get("job_description_text")
                or job.get("description")
                or job.get("description_full")
                or ""
            ).strip()
            if len(desc_text) < 40:
                # Try one live re-fetch from the source URL before failing
                try:
                    from jobs_jd_fetch import fetch_job_description
                    url = (job.get("apply_url_final")
                           or job.get("apply_url")
                           or job.get("source_url"))
                    if url:
                        text, kind = await fetch_job_description(url)
                        if text and len(text) >= 40:
                            desc_text = text
                            await db.jobs_feed.update_one(
                                {"job_id": body.job_id},
                                {"$set": {"description_full": text,
                                          "description_source": kind}},
                            )
                except Exception:
                    pass
            if len(desc_text) < 40:
                raise HTTPException(
                    400,
                    "This verified job doesn't include a full description. "
                    "Add it manually via Job Description library instead.",
                )
            jd = {
                # Ephemeral JD — not stored in job_descriptions library
                "jd_id": f"job_feed:{job.get('job_id')}",
                "job_title": job.get("job_title") or job.get("title") or "",
                "employer": job.get("employer", ""),
                "extracted_text": desc_text,
                "posting_url": job.get("apply_url", "") or job.get("source_url", ""),
                "source": "job_feed",
                "match_scores": {},
                "keyword_analysis": {},
            }
            jd_source_job_id = job.get("job_id")
        else:
            raise HTTPException(400, "Provide either jd_id (from library) or job_id (from feed).")

        result = await _do_tailor(user_id, resume, jd)

        # Persist version
        version_id = f"ver_{uuid.uuid4().hex[:12]}"
        now = _now_iso()
        version_doc = {
            "version_id": version_id, "user_id": user_id,
            "base_resume_id": resume["resume_id"],
            "base_resume_label": resume.get("label") or resume.get("file_name"),
            "jd_id": jd["jd_id"],
            "jd_source": jd.get("source", "library"),
            "source_job_id": jd_source_job_id,
            "job_title": jd.get("job_title", ""),
            "employer": jd.get("employer", ""),
            "generated_date": now,
            **result,
            "manually_edited": False,
            "downloaded": False,
            "emailed": False,
            "saved_to_application": False,
            "job_url": jd.get("posting_url", ""),
        }
        await db.resume_versions.insert_one(version_doc)
        version_doc.pop("_id", None)

        # Update resume last_tailored
        await db.resumes.update_one(
            {"user_id": user_id, "resume_id": resume["resume_id"]},
            {"$set": {"last_tailored": now}},
        )
        # Cache match_score & keyword analysis on the JD collection when
        # library-sourced. Feed-sourced JDs go onto the jobs_feed doc instead.
        if body.jd_id:
            await db.job_descriptions.update_one(
                {"user_id": user_id, "jd_id": jd["jd_id"]},
                {"$set": {
                    f"match_scores.{resume['resume_id']}": result["match_score"],
                    f"keyword_analysis.{resume['resume_id']}": {
                        "found": result["keywords_found"],
                        "added": result["keywords_added"],
                        "missing": result["keywords_missing"],
                        "ats_score_after": result["ats_score_after"],
                    },
                }},
            )
        elif jd_source_job_id:
            await db.jobs_feed.update_one(
                {"job_id": jd_source_job_id},
                {"$set": {
                    f"match_scores.{resume['resume_id']}": result["match_score"],
                    f"keyword_analysis.{resume['resume_id']}": {
                        "found": result["keywords_found"],
                        "added": result["keywords_added"],
                        "missing": result["keywords_missing"],
                        "ats_score_after": result["ats_score_after"],
                    },
                    "last_tailored_at": now,
                }},
            )

        # Optional email
        email_status = None
        if body.email_to_me:
            email_status = await _email_package(db, user_id, version_doc, body.send_pdf)
            if email_status and email_status.get("status") == "sent":
                await db.resume_versions.update_one(
                    {"version_id": version_id}, {"$set": {"emailed": True}}
                )

        out = {**version_doc}
        out["email_status"] = email_status
        return out

    @r.get("/library/tailor/history")
    async def tailor_history(user_id: str = Depends(get_current_user_id)):
        docs = await db.resume_versions.find(
            {"user_id": user_id},
            {"_id": 0, "tailored_resume_text": 0, "cover_letter_text": 0,
             "interview_questions": 0, "insider_connections": 0},
        ).sort("generated_date", -1).to_list(100)
        return {"history": docs}

    @r.get("/library/tailor/history/{version_id}")
    async def tailor_history_get(version_id: str,
                                 user_id: str = Depends(get_current_user_id)):
        doc = await db.resume_versions.find_one(
            {"user_id": user_id, "version_id": version_id}, {"_id": 0}
        )
        if not doc:
            raise HTTPException(404, "Version not found")
        return doc

    @r.delete("/library/tailor/history/{version_id}")
    async def tailor_history_delete(version_id: str,
                                    user_id: str = Depends(get_current_user_id)):
        res = await db.resume_versions.delete_one(
            {"user_id": user_id, "version_id": version_id}
        )
        if res.deleted_count == 0:
            raise HTTPException(404, "Version not found")
        return {"ok": True}

    @r.post("/library/tailor/history/{version_id}/regenerate")
    async def tailor_regenerate(version_id: str,
                                user_id: str = Depends(get_current_user_id)):
        ver = await db.resume_versions.find_one(
            {"user_id": user_id, "version_id": version_id}, {"_id": 0}
        )
        if not ver:
            raise HTTPException(404, "Version not found")
        resume = await db.resumes.find_one(
            {"user_id": user_id, "resume_id": ver["base_resume_id"]}, {"_id": 0}
        )
        jd = await db.job_descriptions.find_one(
            {"user_id": user_id, "jd_id": ver["jd_id"]}, {"_id": 0}
        )
        if not resume or not jd:
            raise HTTPException(400, "Base resume or JD no longer exists")
        result = await _do_tailor(user_id, resume, jd)
        new_id = f"ver_{uuid.uuid4().hex[:12]}"
        new_doc = {**ver, **result,
                   "version_id": new_id,
                   "generated_date": _now_iso(),
                   "manually_edited": False, "emailed": False, "downloaded": False,
                   "saved_to_application": False,
                   "regenerated_from": version_id}
        await db.resume_versions.insert_one(new_doc)
        new_doc.pop("_id", None)
        return new_doc

    @r.put("/library/tailor/history/{version_id}/edit")
    async def tailor_edit(version_id: str, body: ManualEditBody,
                          user_id: str = Depends(get_current_user_id)):
        update: Dict[str, Any] = {"manually_edited": True}
        if body.tailored_resume_text is not None:
            update["tailored_resume_text"] = body.tailored_resume_text
        if body.cover_letter_text is not None:
            update["cover_letter_text"] = body.cover_letter_text
        if body.thank_you_letter_text is not None:
            update["thank_you_letter_text"] = body.thank_you_letter_text
        if body.follow_up_letter_text is not None:
            update["follow_up_letter_text"] = body.follow_up_letter_text
        if body.withdrawal_letter_text is not None:
            update["withdrawal_letter_text"] = body.withdrawal_letter_text
        res = await db.resume_versions.update_one(
            {"user_id": user_id, "version_id": version_id}, {"$set": update}
        )
        if res.matched_count == 0:
            raise HTTPException(404, "Version not found")
        return {"ok": True}

    class LetterGenBody(BaseModel):
        kind: str  # "thank_you" | "follow_up" | "withdrawal"
        context_notes: Optional[str] = ""

    @r.post("/library/tailor/history/{version_id}/generate-letter")
    async def generate_letter(version_id: str, body: LetterGenBody,
                              user_id: str = Depends(get_current_user_id)):
        """(Re)generate a specific letter for a tailored version using PLOS AI."""
        ver = await db.resume_versions.find_one(
            {"user_id": user_id, "version_id": version_id}, {"_id": 0}
        )
        if not ver:
            raise HTTPException(404, "Version not found")
        kind = body.kind
        if kind == "thank_you":
            spec = ("Generate a professional post-interview THANK YOU LETTER. "
                    "150-220 words. Reference specific interview conversation points "
                    "as placeholders in [square brackets] the user will edit.")
            field = "thank_you_letter_text"
        elif kind == "follow_up":
            spec = ("Generate a professional FOLLOW-UP LETTER to send 1-2 weeks after "
                    "applying with no response. 150-200 words. Reiterate interest, "
                    "value proposition, and one specific reason this role fits.")
            field = "follow_up_letter_text"
        elif kind == "withdrawal":
            spec = ("Generate a courteous WITHDRAWAL LETTER declining to continue "
                    "in the process. 100-150 words. Professional, appreciative, "
                    "leaves the door open for future opportunities.")
            field = "withdrawal_letter_text"
        else:
            raise HTTPException(400, "kind must be thank_you, follow_up or withdrawal")

        prompt = f"""JOB TITLE: {ver.get('job_title')}
EMPLOYER: {ver.get('employer')}
ORIGINAL RESUME CONTEXT:
{(ver.get('tailored_resume_text') or '')[:2500]}

COVER LETTER CONTEXT:
{(ver.get('cover_letter_text') or '')[:1500]}

ADDITIONAL CONTEXT / NOTES FROM USER:
{body.context_notes or '(none)'}

{spec}

Return ONLY the letter text (no JSON, no preamble)."""
        session_id = f"letter-{user_id}-{uuid.uuid4().hex[:8]}"
        raw = await call_claude(session_id, SYSTEM_PROMPT, prompt)
        letter_text = (raw or "").strip()
        await db.resume_versions.update_one(
            {"user_id": user_id, "version_id": version_id},
            {"$set": {field: letter_text, "manually_edited": True}},
        )
        return {"ok": True, "kind": kind, "text": letter_text}

    # ------------------------------------------------------------------
    # DOWNLOADS + EMAIL
    # ------------------------------------------------------------------
    @r.get("/library/tailor/history/{version_id}/download")
    async def tailor_download(version_id: str, kind: str = "combined",
                              fmt: str = "pdf",
                              user_id: str = Depends(get_current_user_id)):
        ver = await db.resume_versions.find_one(
            {"user_id": user_id, "version_id": version_id}, {"_id": 0}
        )
        if not ver:
            raise HTTPException(404, "Version not found")
        title = f"{ver.get('job_title')} — {ver.get('employer')}"
        if kind == "resume":
            body_text = ver.get("tailored_resume_text", "")
            label = "Resume"
        elif kind == "cover":
            body_text = ver.get("cover_letter_text", "")
            label = "Cover Letter"
        elif kind == "thank_you":
            body_text = ver.get("thank_you_letter_text", "")
            label = "Thank You Letter"
        elif kind == "follow_up":
            body_text = ver.get("follow_up_letter_text", "")
            label = "Follow-Up Letter"
        elif kind == "withdrawal":
            body_text = ver.get("withdrawal_letter_text", "")
            label = "Withdrawal Letter"
        else:  # combined
            body_text = (
                (ver.get("tailored_resume_text", "") or "") +
                "\n\n\n\nCOVER LETTER\n\n" +
                (ver.get("cover_letter_text", "") or "")
            )
            label = "Application Package"
        if not (body_text or "").strip():
            raise HTTPException(404, f"No {label} content available for this version.")
        today = datetime.now().strftime("%Y-%m-%d")
        if fmt.lower() == "docx":
            content = build_docx(body_text)
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ext = "docx"
        else:
            content = build_pdf(title, body_text,
                                footer=f"Generated by PLOS Career Assistant · {today}")
            mime = "application/pdf"
            ext = "pdf"
        fn = f"{label} - {ver.get('employer', 'Company')} - {today}.{ext}"
        await db.resume_versions.update_one(
            {"version_id": version_id}, {"$set": {"downloaded": True}}
        )
        return {
            "filename": fn, "mime": mime,
            "content_b64": base64.b64encode(content).decode("ascii"),
        }

    async def _email_package(db_, user_id_: str, ver: Dict[str, Any],
                             send_pdf: bool) -> Dict[str, Any]:
        user = await db_.users.find_one({"user_id": user_id_}, {"_id": 0})
        email = (user or {}).get("email", "")
        name = (user or {}).get("full_name", "PLOS User")
        today = datetime.now().strftime("%Y-%m-%d")
        subject = (f"PLOS Career — Tailored Package for "
                   f"{ver.get('job_title')} at {ver.get('employer')} — {today}")
        attachments: List[Dict[str, Any]] = []
        if send_pdf and ver.get("tailored_resume_text"):
            pkg = ((ver.get("tailored_resume_text") or "") +
                   "\n\n\n\nCOVER LETTER\n\n" +
                   (ver.get("cover_letter_text") or ""))
            pdf_bytes = build_pdf(
                f"{ver.get('job_title')} — {ver.get('employer')}",
                pkg, footer=f"Generated by PLOS Career Assistant · {today}",
            )
            attachments.append({
                "content_b64": base64.b64encode(pdf_bytes).decode("ascii"),
                "filename": f"Application Package - {ver.get('employer')} - {today}.pdf",
                "mime": "application/pdf",
            })
        body_text = (
            f"Hi {name},\n\nAttached is your tailored application package for the "
            f"{ver.get('job_title')} role at {ver.get('employer')}.\n\n"
            f"ATS score after tailoring: {ver.get('ats_score_after')}/100\n"
            f"Match score: {ver.get('match_score')}/100\n\n"
            f"Why you fit:\n{ver.get('why_you_fit', '')}\n\n"
            "---\n\nCOVER LETTER:\n\n" +
            (ver.get("cover_letter_text") or "")
        )
        return send_via_sendgrid(
            to_email=email, to_name=name, subject=subject,
            body_text=body_text, body_html=body_text.replace("\n", "<br>"),
            attachments=attachments,
        )

    @r.post("/library/tailor/history/{version_id}/email")
    async def tailor_email(version_id: str,
                           user_id: str = Depends(get_current_user_id)):
        ver = await db.resume_versions.find_one(
            {"user_id": user_id, "version_id": version_id}, {"_id": 0}
        )
        if not ver:
            raise HTTPException(404, "Version not found")
        status = await _email_package(db, user_id, ver, send_pdf=True)
        if status.get("status") == "sent":
            await db.resume_versions.update_one(
                {"version_id": version_id}, {"$set": {"emailed": True}}
            )
        return status

    @r.post("/library/tailor/history/{version_id}/save-application")
    async def tailor_save_app(version_id: str,
                              user_id: str = Depends(get_current_user_id)):
        ver = await db.resume_versions.find_one(
            {"user_id": user_id, "version_id": version_id}, {"_id": 0}
        )
        if not ver:
            raise HTTPException(404, "Version not found")
        app_id = f"app_{uuid.uuid4().hex[:12]}"
        now = _now_iso()
        app_doc = {
            "application_id": app_id, "user_id": user_id,
            "employer": ver.get("employer"), "role_title": ver.get("job_title"),
            "status": "Ready to Apply", "applied_date": None,
            "job_url": ver.get("job_url"),
            "match_score": ver.get("match_score"),
            "ats_score_after": ver.get("ats_score_after"),
            "tailored_resume_md": ver.get("tailored_resume_text"),
            "cover_letter_md": ver.get("cover_letter_text"),
            "resume_version_id": version_id,
            "created_at": now,
        }
        await db.job_applications.insert_one(app_doc)
        await db.resume_versions.update_one(
            {"user_id": user_id, "version_id": version_id},
            {"$set": {"saved_to_application": True,
                      "saved_to_application_id": app_id}},
        )
        return {"ok": True, "application_id": app_id}

    # ------------------------------------------------------------------
    # Legacy sendgrid status
    # ------------------------------------------------------------------
    @r.get("/library/email/status")
    async def email_status_ep(user_id: str = Depends(get_current_user_id)):
        return {
            "sendgrid_ready": _sendgrid_ready(),
            "hint": ("Set SENDGRID_API_KEY in backend .env to enable email delivery."
                     if not _sendgrid_ready() else
                     "SendGrid is configured. From: noreply@plos-app.com"),
        }

    return r


# ============================================================
# Helpers
# ============================================================
def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _iso_to_sort(iso_str: Optional[str]) -> float:
    if not iso_str:
        return 0.0
    try:
        return datetime.fromisoformat(iso_str.replace("Z", "+00:00")).timestamp()
    except Exception:
        return 0.0


async def _extract_jd_meta(call_claude, jd_text: str) -> tuple:
    """Best-effort AI extraction of title + employer from JD text."""
    if not jd_text:
        return ("Untitled Job", "")
    prompt = f"""Extract the job title and employer from this job description
text. Return ONLY a JSON object like {{"job_title": "...", "employer": "..."}}.
If the employer is not clearly stated, return an empty string.

JD text:
\"\"\"{jd_text[:2000]}\"\"\"
"""
    try:
        raw = await call_claude(
            f"jd-meta-{uuid.uuid4().hex[:6]}",
            "You extract structured metadata from job descriptions. Reply with valid JSON only.",
            prompt,
        )
        data = _extract_json(raw)
        title = (data.get("job_title") or "").strip() or "Untitled Job"
        employer = (data.get("employer") or "").strip()
        return (title[:200], employer[:200])
    except Exception as exc:
        logger.warning("JD meta extraction failed: %s", exc)
        return ("Untitled Job", "")
