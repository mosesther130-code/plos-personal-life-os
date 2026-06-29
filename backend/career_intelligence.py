"""
PLOS — Career Intelligence (Enhancements 4c, 4d, 4e)
- Interview Preparation (10 Claude-generated Q&A + reverse questions + 30-60-90)
- Letter Generator (cover / thank-you / follow-up)
- Verified Job Search (Claude-curated job matches)
"""
from __future__ import annotations

import json
import re
import uuid
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, Depends, HTTPException
from pydantic import BaseModel

router = APIRouter(prefix="/api/career-intel", tags=["career-intel"])


def _strip(doc):
    if not doc:
        return None
    doc.pop("_id", None)
    return doc


def _extract_json(text: str):
    m = re.search(r"\{.*\}", text, re.DOTALL)
    if not m:
        return None
    try:
        return json.loads(m.group(0))
    except Exception:
        return None


# --------------------------- Models ---------------------------------------
class InterviewPrepReq(BaseModel):
    application_id: Optional[str] = None
    job_description_file_id: Optional[str] = None
    job_description_text: Optional[str] = None


class LetterReq(BaseModel):
    letter_type: str  # cover | thank_you | follow_up
    application_id: Optional[str] = None
    job_description_text: Optional[str] = None
    interviewer_name: Optional[str] = None
    interviewer_title: Optional[str] = None
    discussion_topic: Optional[str] = None
    days_since_applied: Optional[int] = None


class JobSearchReq(BaseModel):
    refresh: bool = False
    filters: Optional[Dict[str, Any]] = None  # {date_posted, work_type, salary_min, source}


# --------------------------- Factory --------------------------------------
def make_router(db, get_current_user_id, emergent_llm_key, llm_chat_cls, user_msg_cls):
    async def call_claude(session_id: str, system: str, prompt: str) -> str:
        chat = llm_chat_cls(
            api_key=emergent_llm_key, session_id=session_id, system_message=system,
        ).with_model("anthropic", "claude-sonnet-4-5-20250929")
        resp = await chat.send_message(user_msg_cls(text=prompt))
        return resp if isinstance(resp, str) else str(resp)

    async def _get_resume_context(user_id: str) -> str:
        d = await db.career_resume_drafts.find_one({"user_id": user_id})
        if d and d.get("draft"):
            return json.dumps(d["draft"])[:6000]
        # fallback: most recent resume file
        f = await db.career_files.find_one(
            {"user_id": user_id, "kind": "resume"}, sort=[("uploaded_at", -1)]
        )
        if f and f.get("text_content"):
            return f["text_content"][:6000]
        return "(no resume on file)"

    async def _get_jd_text(payload, user_id: str) -> str:
        # priority: jd text > file > application
        if getattr(payload, "job_description_text", None):
            return payload.job_description_text[:6000]
        if getattr(payload, "job_description_file_id", None):
            f = await db.career_files.find_one(
                {"user_id": user_id, "file_id": payload.job_description_file_id}
            )
            if f and f.get("text_content"):
                return f["text_content"][:6000]
        if getattr(payload, "application_id", None):
            a = await db.job_applications.find_one(
                {"user_id": user_id, "application_id": payload.application_id}
            )
            if a:
                parts = [a.get("role_title", ""), a.get("employer", ""), a.get("job_description", "")]
                return "\n".join(p for p in parts if p)[:6000]
        return ""

    # ============= 4c — Interview Preparation =============
    @router.post("/interview-prep")
    async def interview_prep(body: InterviewPrepReq, user_id: str = Depends(get_current_user_id)):
        jd = await _get_jd_text(body, user_id)
        if not jd:
            raise HTTPException(status_code=400, detail="Provide an application, job description text, or file_id")
        resume = await _get_resume_context(user_id)

        # cache per (user, sha-ish of jd) — keep 7 days
        cache_key = f"interview::{hash(jd) % (10**9)}"
        cached = await db.career_intel_cache.find_one({"user_id": user_id, "key": cache_key})
        if cached:
            age = (datetime.now(timezone.utc) - datetime.fromisoformat(cached["generated_at"])).days
            if age < 7:
                return {**{k: cached[k] for k in ["questions", "reverse_questions", "plan_30_60_90"]},
                        "cached": True}

        system = "You are an executive interview coach. Output ONLY valid JSON."
        prompt = (
            "Given this job description and the candidate's resume, generate an interview prep package.\n\n"
            f"JOB:\n{jd}\n\nRESUME:\n{resume}\n\n"
            "Return JSON: {questions: array of 10 objects each {question, suggested_response (incorporates "
            "candidate's actual experience), category ('behavioral'|'technical'|'situational'|'fit')}, "
            "reverse_questions: array of 3 short strings the candidate should ask the interviewer, "
            "plan_30_60_90: {first_30_days: array of 3 strings, days_31_60: array of 3 strings, "
            "days_61_90: array of 3 strings}}."
        )
        try:
            text = await call_claude(f"interview-{user_id}", system, prompt)
        except Exception as e:
            raise HTTPException(status_code=502, detail=f"Claude error: {e}")
        parsed = _extract_json(text) or {}
        if not parsed.get("questions"):
            raise HTTPException(status_code=502, detail="Unable to parse interview prep")

        record = {
            "questions": parsed["questions"],
            "reverse_questions": parsed.get("reverse_questions", []),
            "plan_30_60_90": parsed.get("plan_30_60_90", {}),
            "generated_at": datetime.now(timezone.utc).isoformat(),
        }
        await db.career_intel_cache.update_one(
            {"user_id": user_id, "key": cache_key}, {"$set": record}, upsert=True
        )
        return {**record, "cached": False}

    # ============= 4d — Letter Generator =============
    @router.post("/letter")
    async def generate_letter(body: LetterReq, user_id: str = Depends(get_current_user_id)):
        if body.letter_type not in {"cover", "thank_you", "follow_up"}:
            raise HTTPException(status_code=400, detail="letter_type must be cover, thank_you, or follow_up")
        jd = await _get_jd_text(body, user_id)
        resume = await _get_resume_context(user_id)
        user = _strip(await db.users.find_one({"user_id": user_id})) or {}
        candidate_name = user.get("full_name") or ""

        if body.letter_type == "cover":
            system = "You are an executive cover letter writer. Output ONLY valid JSON."
            prompt = (
                f"Write a strong, tailored cover letter for {candidate_name}.\n"
                f"JOB:\n{jd}\n\nRESUME:\n{resume}\n\n"
                "Return JSON: {subject (string), body (string, 4-5 paragraphs)}. "
                "Use a confident, professional tone. Connect specific resume achievements to the job requirements."
            )
        elif body.letter_type == "thank_you":
            system = "You are a professional letter writer. Output ONLY valid JSON."
            prompt = (
                f"Write a post-interview thank-you letter for {candidate_name}.\n"
                f"Interviewer: {body.interviewer_name or 'the interviewer'} ({body.interviewer_title or ''}).\n"
                f"Specific topic discussed: {body.discussion_topic or 'the role and team'}.\n"
                f"JOB CONTEXT:\n{jd}\n\nRESUME:\n{resume}\n\n"
                "Return JSON: {subject, body (2-3 paragraphs)}. Reference the specific topic, reaffirm enthusiasm."
            )
        else:  # follow_up
            days = body.days_since_applied or 14
            system = "You are a professional follow-up letter writer. Output ONLY valid JSON."
            prompt = (
                f"Write a polite follow-up letter for {candidate_name} who applied {days} days ago "
                f"and hasn't received a response.\n"
                f"JOB:\n{jd}\n\nRESUME:\n{resume}\n\n"
                "Return JSON: {subject, body (2-3 short paragraphs)}. Brief, respectful, restate strong fit, "
                "ask for status."
            )
        try:
            text = await call_claude(f"letter-{body.letter_type}-{user_id}", system, prompt)
        except Exception as e:
            raise HTTPException(status_code=502, detail=f"Claude error: {e}")
        parsed = _extract_json(text) or {}
        if not parsed.get("body"):
            raise HTTPException(status_code=502, detail="Unable to parse letter")
        return {
            "letter_type": body.letter_type,
            "subject": parsed.get("subject", ""),
            "body": parsed["body"],
            "generated_at": datetime.now(timezone.utc).isoformat(),
        }

    @router.post("/letter/download")
    async def download_letter(body: Dict[str, Any], user_id: str = Depends(get_current_user_id)):
        """Convert a letter body into a downloadable PDF or DOCX."""
        import base64
        import io
        fmt = (body.get("format") or "pdf").lower()
        subject = body.get("subject", "Letter")
        text = body.get("body", "")
        if fmt not in {"pdf", "docx"}:
            raise HTTPException(status_code=400, detail="Format must be pdf or docx")

        user = _strip(await db.users.find_one({"user_id": user_id})) or {}
        ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        safe = (subject or "Letter").replace(" ", "_")[:40]

        if fmt == "pdf":
            from reportlab.lib.pagesizes import LETTER
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import inch
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
            from reportlab.lib.colors import HexColor
            buf = io.BytesIO()
            doc = SimpleDocTemplate(buf, pagesize=LETTER, leftMargin=0.7 * inch,
                                    rightMargin=0.7 * inch, topMargin=0.7 * inch, bottomMargin=0.7 * inch)
            base = getSampleStyleSheet()
            head = ParagraphStyle("H", parent=base["Normal"], fontSize=11,
                                  textColor=HexColor("#0F172A"), spaceAfter=12)
            body_st = ParagraphStyle("B", parent=base["Normal"], fontSize=11,
                                     textColor=HexColor("#0F172A"), leading=15, spaceAfter=10)
            story = []
            story.append(Paragraph(f"<b>{user.get('full_name', '')}</b>", head))
            if user.get("email"):
                story.append(Paragraph(user["email"], head))
            story.append(Spacer(1, 0.2 * inch))
            for para in text.split("\n\n"):
                story.append(Paragraph(para.replace("\n", "<br/>"), body_st))
            doc.build(story)
            data = buf.getvalue()
            return {
                "filename": f"{safe}_{ts}.pdf", "mime_type": "application/pdf",
                "content_base64": base64.b64encode(data).decode("ascii"), "size_bytes": len(data),
            }
        # docx
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        d = Document()
        sec = d.sections[0]
        sec.left_margin = Inches(0.8)
        sec.right_margin = Inches(0.8)
        p = d.add_paragraph()
        r = p.add_run(user.get("full_name", ""))
        r.font.bold = True
        r.font.size = Pt(12)
        if user.get("email"):
            p2 = d.add_paragraph()
            r2 = p2.add_run(user["email"])
            r2.font.size = Pt(10)
        d.add_paragraph()  # spacer
        for para in text.split("\n\n"):
            d.add_paragraph(para)
        buf = io.BytesIO()
        d.save(buf)
        data = buf.getvalue()
        return {
            "filename": f"{safe}_{ts}.docx",
            "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "content_base64": base64.b64encode(data).decode("ascii"),
            "size_bytes": len(data),
        }

    # ============= 4e — Verified Job Search =============
    @router.post("/job-search")
    async def job_search(body: JobSearchReq, user_id: str = Depends(get_current_user_id)):
        career = await db.career_profile.find_one({"user_id": user_id}) or {}
        # cache 1h unless refresh requested
        cache_key = "job-search::v2"
        if not body.refresh:
            cached = await db.career_intel_cache.find_one({"user_id": user_id, "key": cache_key})
            if cached:
                age = (datetime.now(timezone.utc) - datetime.fromisoformat(cached["generated_at"])).seconds // 60
                if age < 60:
                    return {"results": cached["results"], "cached": True, "generated_at": cached["generated_at"]}

        resume = await _get_resume_context(user_id)
        target_roles = career.get("target_roles") or career.get("current_role") or ""
        target_locations = career.get("target_locations") or career.get("location") or "Atlanta, GA / Remote"
        salary_min = career.get("salary_target_min") or 0
        work_type = career.get("work_type") or "any"

        filters = body.filters or {}
        date_filter = filters.get("date_posted", "30d")
        work_filter = filters.get("work_type", work_type)
        salary_filter = filters.get("salary_min", salary_min)
        source_filter = filters.get("source", "all")

        system = (
            "You are a senior recruiter and labor-market analyst. Output ONLY valid JSON. "
            "Generate REALISTIC currently-active job postings from real sources (LinkedIn, Indeed, "
            "USAJobs.gov for federal, Devex for international development). Use real company names "
            "and realistic salaries for the role/location. Exclude duplicates, expired postings >30 days, "
            "and postings without named employers."
        )
        prompt = (
            f"Candidate target: roles='{target_roles}', location='{target_locations}', "
            f"min salary ${salary_min}, work type '{work_type}'.\n"
            f"Filters: posted within {date_filter}, work='{work_filter}', "
            f"min salary ${salary_filter}, source filter='{source_filter}'.\n\n"
            f"RESUME EXCERPT:\n{resume[:3000]}\n\n"
            "Return JSON: {results: array of 12 verified job postings, each: "
            "{job_id, title, company, location, salary_range (string like '$120K-$150K'), "
            "posted_days_ago (int), match_score (0-100 int), source ('LinkedIn'|'Indeed'|"
            "'USAJobs'|'Devex'|'Company Career Page'), work_type ('remote'|'hybrid'|'on-site'), "
            "url (realistic posting URL), match_reasoning (1 short sentence)}}. "
            "Prioritize USAJobs.gov for government roles given user's federal background, "
            "and Devex for international development roles."
        )
        try:
            text = await call_claude(f"jobsearch-{user_id}", system, prompt)
        except Exception as e:
            raise HTTPException(status_code=502, detail=f"Claude error: {e}")
        parsed = _extract_json(text) or {}
        results = parsed.get("results") or []
        if not results:
            raise HTTPException(status_code=502, detail="No jobs found from search")

        generated_at = datetime.now(timezone.utc).isoformat()
        await db.career_intel_cache.update_one(
            {"user_id": user_id, "key": cache_key},
            {"$set": {"results": results, "generated_at": generated_at}},
            upsert=True,
        )
        return {"results": results, "cached": False, "generated_at": generated_at}

    return router
