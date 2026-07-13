"""PLOS Career — AI Resume Tailor pipeline tests (iteration 27).

Tests the resume vault CRUD + PLOS AI Sonnet 4.5 tailor endpoint + PDF gen +
thankyou/followup/save-application + SendGrid deferred behaviour.
"""
import base64
import io
import os
import re
import time

import pytest
import requests

BASE_URL = os.environ.get("BACKEND_URL", "http://localhost:8001").rstrip("/")
EMAIL = "test1@plos.app"
PASSWORD = "test123"

# ---------------- Fixtures ----------------
@pytest.fixture(scope="module")
def token():
    r = requests.post(
        f"{BASE_URL}/api/auth/login",
        json={"email": EMAIL, "password": PASSWORD},
        timeout=30,
    )
    assert r.status_code == 200, f"Login failed: {r.status_code} {r.text}"
    return r.json()["token"]


@pytest.fixture(scope="module")
def auth(token):
    return {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}


# Real 800+ char finance resume for paste tests / tailor
RESUME_TEXT = """Moses Ndifon
Atlanta, GA | moses.ndifon@example.com | (404) 555-0134 | linkedin.com/in/mosesndifon

PROFESSIONAL SUMMARY
Finance and administration leader with 12+ years of experience spanning USAID foreign-assistance
portfolio management, multilateral development finance, and higher-education administration.
MBA-credentialed accountant skilled in IPSAS-adjacent reporting, grants management, portfolio oversight,
budget forecasting, and financial control for multi-country programs across Asia-Pacific. Fluent in
GAAP and federal Cost Principles (2 CFR 200). Proven partner to ADB, World Bank, and USAID mission
teams delivering audit-ready results.

CORE COMPETENCIES
Financial Control | Grants Management | IPSAS & GAAP | Portfolio Oversight | Budget Forecasting
Multi-Country Operations | USAID Cost Principles | Internal Controls | Financial Reporting
Stakeholder Management | Academic Administration | ADB Programs

PROFESSIONAL EXPERIENCE
Department Coordinator — Georgia State University, Perimeter College (2021–Present)
- Manage $4.2M departmental operating budget, monthly variance reporting to the Dean.
- Coordinate cross-college financial reporting and student-fee reconciliation processes.
- Streamlined purchase-order workflow, reducing cycle time 38%.

Deputy Controller — USAID / Cardno Emerging Markets (2015–2021)
- Oversaw a $180M multi-country foreign-assistance portfolio across the Asia-Pacific region.
- Managed grants and cooperative agreements, ensured 2 CFR 200 compliance, zero audit findings 3 yrs.
- Built forecasting model that improved obligation accuracy by 22%.

Senior Accountant — Various (2010–2015)
- Prepared consolidated financial statements, monthly close, IPSAS-adjacent donor reporting.

EDUCATION
MBA, University of West Georgia
BBA, Accounting — Morehead State University

CERTIFICATIONS
CGFM (Certified Government Financial Manager) — candidate
"""
assert len(RESUME_TEXT) > 800

# Real ~1200 char JD for the tailor test
JD_TEXT = """The Asian Development Bank (ADB) is seeking a Financial Control Specialist to join its
Controller's Department in Manila. The role provides financial control, grants management, and
portfolio oversight for multilateral development operations across the Asia-Pacific region.

Key responsibilities:
- Lead financial control activities for ADB grants and technical assistance programs,
  ensuring compliance with IPSAS and internal accounting policies.
- Support portfolio oversight of loans and grants across multi-country programs, including
  monitoring drawdowns, budget forecasting, and quarterly financial reporting.
- Coordinate closely with USAID and other bilateral donors on co-financed operations; prior
  USAID experience strongly preferred.
- Design and enforce internal controls, review disbursement requests, and lead year-end audit
  readiness for the multilateral finance portfolio.
- Prepare consolidated financial reports and management commentary for senior leadership.

Required qualifications:
- MBA required; professional accounting credential (CPA, CGFM, ACCA) strongly preferred.
- 10+ years of progressive experience in financial control, grants management, or multilateral
  development finance.
- Deep familiarity with IPSAS, GAAP, and federal Cost Principles.
- Multi-country experience across Asia-Pacific, ADB, or comparable multilateral institutions.
- Strong Excel modelling; ERP experience (Oracle, SAP) preferred.

ATS keywords: financial control, IPSAS, grants management, portfolio oversight, multi-country
experience, MBA required, 10+ years experience, financial reporting, USAID experience preferred,
ADB grants, budget forecasting, multilateral finance.
"""
assert len(JD_TEXT) > 1200


def _wipe_resumes(auth):
    r = requests.get(f"{BASE_URL}/api/career/resumes", headers=auth, timeout=30)
    if r.status_code == 200:
        for item in r.json().get("resumes", []):
            requests.delete(
                f"{BASE_URL}/api/career/resumes/{item['resume_id']}",
                headers=auth, timeout=30,
            )


# ---------------- Resume Vault CRUD ----------------
class TestResumeVault:
    def test_00_wipe(self, auth):
        _wipe_resumes(auth)
        r = requests.get(f"{BASE_URL}/api/career/resumes", headers=auth, timeout=30)
        assert r.status_code == 200
        assert r.json().get("resumes") == []

    def test_01_paste_first_becomes_default(self, auth):
        r = requests.post(
            f"{BASE_URL}/api/career/resumes", headers=auth,
            json={"name": "TEST_Master", "file_type": "paste", "text": RESUME_TEXT},
            timeout=30,
        )
        assert r.status_code == 201, r.text
        data = r.json()
        assert data["resume_id"].startswith("res_")
        assert data["is_default"] is True
        assert "content_b64" not in data
        pytest.resume_id_1 = data["resume_id"]

    def test_02_second_does_not_auto_default(self, auth):
        r = requests.post(
            f"{BASE_URL}/api/career/resumes", headers=auth,
            json={"name": "TEST_Second", "file_type": "paste", "text": RESUME_TEXT + "\nExtra."},
            timeout=30,
        )
        assert r.status_code == 201
        data = r.json()
        assert data["is_default"] is False
        pytest.resume_id_2 = data["resume_id"]

    def test_03_list_strips_content_b64(self, auth):
        r = requests.get(f"{BASE_URL}/api/career/resumes", headers=auth, timeout=30)
        assert r.status_code == 200
        items = r.json()["resumes"]
        assert len(items) >= 2
        for it in items:
            assert "content_b64" not in it
        ids = [i["resume_id"] for i in items]
        assert pytest.resume_id_1 in ids and pytest.resume_id_2 in ids

    def test_04_set_default_swaps(self, auth):
        r = requests.put(
            f"{BASE_URL}/api/career/resumes/{pytest.resume_id_2}",
            headers=auth, json={"is_default": True}, timeout=30,
        )
        assert r.status_code == 200
        assert "is_default" in r.json().get("updated", [])
        # Verify swap
        r2 = requests.get(f"{BASE_URL}/api/career/resumes", headers=auth, timeout=30)
        by_id = {i["resume_id"]: i for i in r2.json()["resumes"]}
        assert by_id[pytest.resume_id_2]["is_default"] is True
        assert by_id[pytest.resume_id_1]["is_default"] is False

    def test_05_rename(self, auth):
        r = requests.put(
            f"{BASE_URL}/api/career/resumes/{pytest.resume_id_1}",
            headers=auth, json={"name": "Renamed"}, timeout=30,
        )
        assert r.status_code == 200
        assert r.json().get("updated") == ["name"]

    def test_06_delete_promotes_newest(self, auth):
        # Delete the current default (resume_id_2 - the newer one), newest remaining should get promoted
        r = requests.delete(
            f"{BASE_URL}/api/career/resumes/{pytest.resume_id_2}",
            headers=auth, timeout=30,
        )
        assert r.status_code == 200
        r2 = requests.get(f"{BASE_URL}/api/career/resumes", headers=auth, timeout=30)
        items = r2.json()["resumes"]
        assert any(i["is_default"] for i in items), "A remaining resume should be promoted to default"

    def test_07_pdf_upload(self, auth):
        # Build a real small PDF with fpdf2
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        pdf.multi_cell(0, 8, "TEST PDF RESUME\nMoses Ndifon\nFinance professional.\nMBA, University of West Georgia.")
        out = pdf.output(dest="S")
        pdf_bytes = out.encode("latin-1") if isinstance(out, str) else bytes(out)
        b64 = base64.b64encode(pdf_bytes).decode("ascii")
        r = requests.post(
            f"{BASE_URL}/api/career/resumes", headers=auth,
            json={"name": "TEST_PDF", "file_type": "pdf", "content_b64": b64},
            timeout=60,
        )
        assert r.status_code == 201, r.text
        d = r.json()
        assert d["text"], "pdfplumber extraction should populate text"
        assert d["size_bytes"] > 0
        assert "Moses Ndifon" in d["text"] or "TEST PDF" in d["text"].upper()
        pytest.resume_pdf_id = d["resume_id"]

    def test_08_docx_upload(self, auth):
        from docx import Document
        doc = Document()
        doc.add_heading("Moses Ndifon", level=0)
        doc.add_paragraph("Finance professional with USAID and ADB experience.")
        doc.add_paragraph("MBA, University of West Georgia.")
        buf = io.BytesIO()
        doc.save(buf)
        b64 = base64.b64encode(buf.getvalue()).decode("ascii")
        r = requests.post(
            f"{BASE_URL}/api/career/resumes", headers=auth,
            json={"name": "TEST_DOCX", "file_type": "docx", "content_b64": b64},
            timeout=60,
        )
        assert r.status_code == 201, r.text
        d = r.json()
        assert d["text"] and "Moses" in d["text"]
        pytest.resume_docx_id = d["resume_id"]

    def test_09_paste_empty_400(self, auth):
        r = requests.post(
            f"{BASE_URL}/api/career/resumes", headers=auth,
            json={"name": "empty", "file_type": "paste", "text": ""}, timeout=30,
        )
        assert r.status_code == 400

    def test_10_put_bogus_404(self, auth):
        r = requests.put(
            f"{BASE_URL}/api/career/resumes/res_doesnotexist",
            headers=auth, json={"name": "x"}, timeout=30,
        )
        assert r.status_code == 404


# ---------------- Tailor endpoint ----------------
class TestTailor:
    def _ensure_default_paste_resume(self, auth):
        # Ensure the default resume has substantial text (paste type)
        _wipe_resumes(auth)
        r = requests.post(
            f"{BASE_URL}/api/career/resumes", headers=auth,
            json={"name": "TEST_Tailor_Master", "file_type": "paste", "text": RESUME_TEXT},
            timeout=30,
        )
        assert r.status_code == 201

    def test_11_tailor_full_bundle(self, auth):
        self._ensure_default_paste_resume(auth)
        payload = {
            "job_title": "Financial Control Specialist",
            "company": "Asian Development Bank",
            "job_description": JD_TEXT,
            "job_url": "https://www.adb.org/work-with-us/careers",
            "tailor_resume": True,
            "generate_cover_letter": True,
            "generate_interview_questions": True,
            "email_to_me": False,
        }
        t0 = time.time()
        r = requests.post(f"{BASE_URL}/api/career/tailor", headers=auth,
                          json=payload, timeout=120)
        # Retry ONCE on 502 (PLOS AI parseable JSON)
        if r.status_code == 502:
            time.sleep(2)
            r = requests.post(f"{BASE_URL}/api/career/tailor", headers=auth,
                              json=payload, timeout=120)
        elapsed = time.time() - t0
        assert r.status_code == 200, f"Tailor failed: {r.status_code} {r.text[:400]}"
        assert elapsed < 120, f"Too slow: {elapsed}s"

        data = r.json()
        pytest.version_id = data["version_id"]
        assert pytest.version_id.startswith("ver_")
        ats = data["ats_score"]
        assert isinstance(ats, int) and 0 <= ats <= 100
        pytest.ats_score = ats

        km = data["keywords_matched"]
        assert isinstance(km, list) and len(km) >= 5
        assert all(isinstance(k, str) and k.strip() for k in km)
        assert isinstance(data["keywords_missing"], list)

        resume_md = data["tailored_resume_md"]
        assert resume_md and len(resume_md) > 400
        rl = resume_md.lower()
        assert "experience" in rl and "education" in rl

        cover = data["cover_letter_md"]
        assert cover
        wc = len(re.findall(r"\S+", cover))
        pytest.cover_wc = wc
        assert 300 <= wc <= 500, f"Cover word count {wc} not in 300-500"

        qs = data["interview_questions"]
        assert isinstance(qs, list) and len(qs) == 10

        assert data["summary"]
        assert data["email_status"] is None

        # Hallucination check
        matched_present = sum(1 for k in km if k.lower() in rl)
        pytest.matched_present = matched_present
        pytest.km_total = len(km)
        assert matched_present >= max(1, int(len(km) * 0.5)), (
            f"Only {matched_present}/{len(km)} keywords_matched appear in tailored resume"
        )
        print(
            f"\n[TAILOR RESULT] ATS={ats}  keywords_matched={len(km)} "
            f"present_in_resume={matched_present}  cover_wc={wc} elapsed={elapsed:.1f}s"
        )

    def test_12_list_versions_strips_heavy(self, auth):
        r = requests.get(f"{BASE_URL}/api/career/tailor/versions",
                         headers=auth, timeout=30)
        assert r.status_code == 200
        vs = r.json()["versions"]
        assert any(v["version_id"] == pytest.version_id for v in vs)
        for v in vs:
            assert "tailored_resume_md" not in v
            assert "cover_letter_md" not in v
            assert "interview_questions" not in v

    def test_13_get_version_full(self, auth):
        r = requests.get(
            f"{BASE_URL}/api/career/tailor/versions/{pytest.version_id}",
            headers=auth, timeout=30,
        )
        assert r.status_code == 200
        d = r.json()
        assert d["tailored_resume_md"] and d["cover_letter_md"]
        assert len(d["interview_questions"]) == 10

    def test_14_download_resume_pdf(self, auth):
        r = requests.get(
            f"{BASE_URL}/api/career/tailor/versions/{pytest.version_id}/download",
            headers=auth, params={"kind": "resume"}, timeout=30,
        )
        assert r.status_code == 200, r.text
        d = r.json()
        assert d["mime"] == "application/pdf"
        raw = base64.b64decode(d["content_b64"])
        assert raw.startswith(b"%PDF-"), "PDF magic bytes missing"
        assert len(raw) >= 2048, f"PDF too small: {len(raw)}B"
        pytest.pdf_resume_bytes = len(raw)

    def test_15_download_cover_pdf(self, auth):
        r = requests.get(
            f"{BASE_URL}/api/career/tailor/versions/{pytest.version_id}/download",
            headers=auth, params={"kind": "cover"}, timeout=30,
        )
        assert r.status_code == 200
        raw = base64.b64decode(r.json()["content_b64"])
        assert raw.startswith(b"%PDF-")
        pytest.pdf_cover_bytes = len(raw)

    def test_16_download_thankyou_before_gen_404(self, auth):
        r = requests.get(
            f"{BASE_URL}/api/career/tailor/versions/{pytest.version_id}/download",
            headers=auth, params={"kind": "thankyou"}, timeout=30,
        )
        assert r.status_code == 404

    def test_17_thankyou_generation(self, auth):
        r = requests.post(
            f"{BASE_URL}/api/career/tailor/thankyou", headers=auth,
            json={
                "version_id": pytest.version_id,
                "interviewer_name": "Jane Doe",
                "topic_discussed": "Their new regional finance transformation initiative",
            }, timeout=90,
        )
        assert r.status_code == 200, r.text
        letter = r.json()["thank_you_letter_md"]
        assert letter
        assert "Jane Doe" in letter
        assert re.search(r"regional finance|finance transformation", letter, re.I)

    def test_18_download_thankyou_after_gen(self, auth):
        r = requests.get(
            f"{BASE_URL}/api/career/tailor/versions/{pytest.version_id}/download",
            headers=auth, params={"kind": "thankyou"}, timeout=30,
        )
        assert r.status_code == 200
        raw = base64.b64decode(r.json()["content_b64"])
        assert raw.startswith(b"%PDF-")

    def test_19_followup_generation(self, auth):
        r = requests.post(
            f"{BASE_URL}/api/career/tailor/followup", headers=auth,
            json={"version_id": pytest.version_id, "days_since_applied": 10},
            timeout=90,
        )
        assert r.status_code == 200, r.text
        assert r.json()["follow_up_letter_md"]

    def test_20_save_application(self, auth):
        r = requests.post(
            f"{BASE_URL}/api/career/tailor/save-application", headers=auth,
            json={"version_id": pytest.version_id}, timeout=30,
        )
        assert r.status_code == 200, r.text
        app_id = r.json()["application_id"]
        assert app_id.startswith("app_")
        # Verify persistence
        r2 = requests.get(f"{BASE_URL}/api/career/applications",
                          headers=auth, timeout=30)
        if r2.status_code == 200:
            apps = r2.json() if isinstance(r2.json(), list) else r2.json().get("applications", [])
            found = [a for a in apps if a.get("application_id") == app_id]
            if found:
                assert found[0].get("status") == "Applied"

    def test_21_email_status_placeholder(self, auth):
        r = requests.get(
            f"{BASE_URL}/api/career/tailor/email/status",
            headers=auth, timeout=30,
        )
        assert r.status_code == 200
        d = r.json()
        assert d["sendgrid_ready"] is False
        assert d.get("hint")

    def test_22_tailor_with_email_deferred(self, auth):
        payload = {
            "job_title": "Financial Control Specialist",
            "company": "Asian Development Bank",
            "job_description": JD_TEXT,
            "tailor_resume": True,
            "generate_cover_letter": False,
            "generate_interview_questions": False,
            "email_to_me": True,
        }
        r = requests.post(f"{BASE_URL}/api/career/tailor",
                          headers=auth, json=payload, timeout=120)
        if r.status_code == 502:
            time.sleep(2)
            r = requests.post(f"{BASE_URL}/api/career/tailor",
                              headers=auth, json=payload, timeout=120)
        assert r.status_code == 200, r.text
        es = r.json().get("email_status")
        assert es is not None
        # Expected behaviour: deferred because SENDGRID_API_KEY=placeholder
        assert es.get("status") == "deferred"
        assert es.get("reason") == "sendgrid_key_missing"

    def test_23_no_resume_400(self, auth):
        _wipe_resumes(auth)
        r = requests.post(
            f"{BASE_URL}/api/career/tailor", headers=auth,
            json={
                "job_title": "X", "company": "Y", "job_description": JD_TEXT,
            }, timeout=30,
        )
        assert r.status_code == 400
        assert "resume" in r.text.lower()

    def test_24_empty_jd_422_or_400(self, auth):
        # Restore a resume so we test JD validation specifically
        requests.post(
            f"{BASE_URL}/api/career/resumes", headers=auth,
            json={"name": "TEST_edge", "file_type": "paste", "text": RESUME_TEXT},
            timeout=30,
        )
        r = requests.post(
            f"{BASE_URL}/api/career/tailor", headers=auth,
            json={
                "job_title": "X", "company": "Y", "job_description": "short",
            }, timeout=30,
        )
        assert r.status_code in (400, 422)

    def test_25_delete_version(self, auth):
        r = requests.delete(
            f"{BASE_URL}/api/career/tailor/versions/{pytest.version_id}",
            headers=auth, timeout=30,
        )
        assert r.status_code == 200
        r2 = requests.get(
            f"{BASE_URL}/api/career/tailor/versions/{pytest.version_id}",
            headers=auth, timeout=30,
        )
        assert r2.status_code == 404


def test_zz_print_summary():
    """Print aggregated metrics for the test report."""
    ats = getattr(pytest, "ats_score", None)
    print(
        f"\n=== SUMMARY ===\n"
        f"ATS score: {ats}\n"
        f"keywords_matched present in resume: "
        f"{getattr(pytest, 'matched_present', None)}/{getattr(pytest, 'km_total', None)}\n"
        f"Cover letter word count: {getattr(pytest, 'cover_wc', None)}\n"
        f"PDF resume bytes: {getattr(pytest, 'pdf_resume_bytes', None)}\n"
        f"PDF cover bytes: {getattr(pytest, 'pdf_cover_bytes', None)}\n"
    )
