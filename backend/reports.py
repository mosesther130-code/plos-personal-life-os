"""
PLOS Financial Reports
Generates Statement (income/expense), Snapshot, and Detailed reports as PDF, DOCX, or CSV.
"""
from __future__ import annotations

import csv
import io
from datetime import datetime
from typing import Any, Dict, List, Tuple

from reportlab.lib import colors as rl_colors
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

# --------------------------- Brand palette --------------------------------
BRAND_BLUE = HexColor("#1E40AF")
BRAND_DARK = HexColor("#0F172A")
BRAND_GREY = HexColor("#64748B")
BRAND_LIGHT = HexColor("#F1F5F9")
BRAND_GREEN = HexColor("#10B981")
BRAND_RED = HexColor("#EF4444")

DOCX_BLUE = RGBColor(0x1E, 0x40, 0xAF)
DOCX_GREY = RGBColor(0x64, 0x75, 0x8B)
DOCX_DARK = RGBColor(0x0F, 0x17, 0x23)


# --------------------------- Helpers --------------------------------------
def _months_between(start: str, end: str) -> int:
    a = datetime.strptime(start, "%Y-%m-%d")
    b = datetime.strptime(end, "%Y-%m-%d")
    months = (b.year - a.year) * 12 + (b.month - a.month) + 1
    return max(1, months)


def _fmt_money(n: float) -> str:
    return f"${n:,.2f}"


def _fmt_pct(n: float) -> str:
    return f"{n:.1f}%"


def _date_range_label(start: str, end: str) -> str:
    a = datetime.strptime(start, "%Y-%m-%d")
    b = datetime.strptime(end, "%Y-%m-%d")
    return f"{a.strftime('%b %d, %Y')} – {b.strftime('%b %d, %Y')}"


def _aggregate(income: list, expenses: list, debts: list, assets: list, months: int) -> Dict[str, Any]:
    monthly_income = sum(float(i.get("net_monthly") or 0) for i in income if i.get("is_active", True))
    monthly_expense = sum(float(e.get("monthly_amount") or 0) for e in expenses)
    monthly_surplus = monthly_income - monthly_expense
    period_income = monthly_income * months
    period_expense = monthly_expense * months
    period_surplus = monthly_surplus * months

    by_category: Dict[str, float] = {}
    for e in expenses:
        c = e.get("category") or "Other"
        by_category[c] = by_category.get(c, 0.0) + float(e.get("monthly_amount") or 0)

    total_debt = sum(float(d.get("balance") or 0) for d in debts)
    total_assets = sum(float(a.get("current_value") or 0) for a in assets)
    net_worth = total_assets - total_debt

    return {
        "monthly_income": monthly_income,
        "monthly_expense": monthly_expense,
        "monthly_surplus": monthly_surplus,
        "period_income": period_income,
        "period_expense": period_expense,
        "period_surplus": period_surplus,
        "by_category": by_category,
        "total_debt": total_debt,
        "total_assets": total_assets,
        "net_worth": net_worth,
    }


# ============================ Public API ===================================
def generate_report(
    report_type: str,
    fmt: str,
    start_date: str,
    end_date: str,
    user: Dict[str, Any],
    income: List[Dict],
    expenses: List[Dict],
    debts: List[Dict],
    assets: List[Dict],
    investments: List[Dict],
) -> Tuple[str, str, bytes]:
    """
    Returns (filename, mime_type, file_bytes).
    report_type: statement_income | statement_expenses | snapshot | detailed
    fmt: pdf | docx | csv (csv only valid for statement_*)
    """
    if report_type not in {"statement_income", "statement_expenses", "snapshot", "detailed"}:
        raise ValueError(f"Unknown report_type: {report_type}")
    if fmt not in {"pdf", "docx", "csv"}:
        raise ValueError(f"Unknown format: {fmt}")
    if fmt == "csv" and not report_type.startswith("statement_"):
        raise ValueError("CSV format is only available for statement reports")

    months = _months_between(start_date, end_date)
    safe_name = (user.get("full_name") or "user").replace(" ", "_").replace("/", "_")
    ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    base_filename = f"PLOS_{report_type}_{safe_name}_{ts}"

    if fmt == "csv":
        data, mime = _gen_csv(report_type, start_date, end_date, months, income, expenses)
        return f"{base_filename}.csv", mime, data
    if fmt == "pdf":
        data = _gen_pdf(report_type, start_date, end_date, months, user, income, expenses, debts, assets, investments)
        return f"{base_filename}.pdf", "application/pdf", data
    # docx
    data = _gen_docx(report_type, start_date, end_date, months, user, income, expenses, debts, assets, investments)
    return (
        f"{base_filename}.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        data,
    )


# ============================ CSV =========================================
def _gen_csv(report_type: str, start: str, end: str, months: int, income: list, expenses: list) -> Tuple[bytes, str]:
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(["PLOS Financial Statement", _date_range_label(start, end), f"Months in period: {months}"])
    w.writerow([])
    if report_type == "statement_income":
        w.writerow(["Source", "Type", "Status", "Monthly Net (USD)", "Months in Period", "Period Total (USD)"])
        grand = 0.0
        for i in income:
            net = float(i.get("net_monthly") or 0)
            active = i.get("is_active", True)
            period_total = net * months if active else 0
            grand += period_total
            w.writerow([
                i.get("source_name", ""),
                i.get("type", ""),
                "Active" if active else "Inactive",
                f"{net:.2f}",
                months if active else 0,
                f"{period_total:.2f}",
            ])
        w.writerow([])
        w.writerow(["", "", "", "", "Total:", f"{grand:.2f}"])
    else:
        w.writerow(["Vendor", "Category", "Monthly Amount (USD)", "Months in Period", "Period Total (USD)", "Auto-Pay", "Due Day"])
        grand = 0.0
        for e in expenses:
            amt = float(e.get("monthly_amount") or 0)
            period_total = amt * months
            grand += period_total
            w.writerow([
                e.get("vendor", ""),
                e.get("category", ""),
                f"{amt:.2f}",
                months,
                f"{period_total:.2f}",
                "Yes" if e.get("auto_pay") else "No",
                e.get("due_day_of_month", ""),
            ])
        w.writerow([])
        w.writerow(["", "", "", "", "Total:", f"{grand:.2f}"])

    return buf.getvalue().encode("utf-8"), "text/csv"


# ============================ PDF =========================================
def _pdf_styles():
    base = getSampleStyleSheet()
    title = ParagraphStyle("Title", parent=base["Title"], fontSize=22, textColor=BRAND_DARK,
                           spaceAfter=4, leading=26, fontName="Helvetica-Bold")
    subtitle = ParagraphStyle("Subtitle", parent=base["Normal"], fontSize=10,
                              textColor=BRAND_GREY, spaceAfter=18)
    section = ParagraphStyle("Section", parent=base["Heading2"], fontSize=13, textColor=BRAND_BLUE,
                             spaceBefore=14, spaceAfter=8, fontName="Helvetica-Bold")
    body = ParagraphStyle("Body", parent=base["Normal"], fontSize=10, textColor=BRAND_DARK, leading=14)
    small = ParagraphStyle("Small", parent=base["Normal"], fontSize=8, textColor=BRAND_GREY, leading=10)
    return title, subtitle, section, body, small


def _pdf_header(story, user, start, end, title_text, subtitle_text):
    title, subtitle, *_ = _pdf_styles()
    story.append(Paragraph("PLOS — Personal Life Operating System", subtitle))
    story.append(Paragraph(title_text, title))
    user_line = f"Prepared for: <b>{user.get('full_name', 'User')}</b> &nbsp;|&nbsp; {user.get('email', '')}"
    story.append(Paragraph(user_line, subtitle))
    story.append(Paragraph(f"Period: <b>{_date_range_label(start, end)}</b> &nbsp;|&nbsp; Generated: {datetime.utcnow().strftime('%b %d, %Y %H:%M UTC')}", subtitle))
    if subtitle_text:
        story.append(Paragraph(subtitle_text, subtitle))


def _pdf_footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(BRAND_GREY)
    page = canvas.getPageNumber()
    text = f"PLOS Financial Report  ·  Page {page}  ·  Confidential"
    canvas.drawCentredString(LETTER[0] / 2, 0.4 * inch, text)
    canvas.setStrokeColor(BRAND_BLUE)
    canvas.setLineWidth(2)
    canvas.line(0.6 * inch, 0.6 * inch, LETTER[0] - 0.6 * inch, 0.6 * inch)
    canvas.restoreState()


def _table_style(header=True):
    cmd = [
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("TEXTCOLOR", (0, 0), (-1, -1), BRAND_DARK),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [rl_colors.white, BRAND_LIGHT]),
        ("GRID", (0, 0), (-1, -1), 0.25, BRAND_GREY),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]
    if header:
        cmd += [
            ("BACKGROUND", (0, 0), (-1, 0), BRAND_BLUE),
            ("TEXTCOLOR", (0, 0), (-1, 0), rl_colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 9),
        ]
    return TableStyle(cmd)


def _bar_chart_table(income: float, expense: float) -> Table:
    """Simple two-row bar chart using table cells."""
    mx = max(income, expense, 1)
    inc_w = (income / mx) * 4.0
    exp_w = (expense / mx) * 4.0
    data = [
        ["Income", Table([[""]], colWidths=[inc_w * inch], rowHeights=[0.25 * inch],
                         style=TableStyle([("BACKGROUND", (0, 0), (-1, -1), BRAND_GREEN)])),
         _fmt_money(income)],
        ["Expenses", Table([[""]], colWidths=[exp_w * inch], rowHeights=[0.25 * inch],
                           style=TableStyle([("BACKGROUND", (0, 0), (-1, -1), BRAND_RED)])),
         _fmt_money(expense)],
    ]
    t = Table(data, colWidths=[1.0 * inch, 4.2 * inch, 1.4 * inch])
    t.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("TEXTCOLOR", (0, 0), (0, -1), BRAND_DARK),
        ("TEXTCOLOR", (2, 0), (2, -1), BRAND_DARK),
        ("ALIGN", (2, 0), (2, -1), "RIGHT"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
    ]))
    return t


def _gen_pdf(report_type, start, end, months, user, income, expenses, debts, assets, investments) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER,
                            leftMargin=0.6 * inch, rightMargin=0.6 * inch,
                            topMargin=0.6 * inch, bottomMargin=0.8 * inch,
                            title="PLOS Financial Report")
    story = []
    _, _, section, body, _ = _pdf_styles()
    agg = _aggregate(income, expenses, debts, assets, months)

    if report_type == "statement_income":
        _pdf_header(story, user, start, end, "Income Statement", "All income sources for the selected period.")
        story.append(Paragraph("Income Sources", section))
        rows = [["Source", "Type", "Status", "Monthly Net", "× Months", "Period Total"]]
        grand = 0.0
        for i in income:
            net = float(i.get("net_monthly") or 0)
            active = i.get("is_active", True)
            total = net * months if active else 0
            grand += total
            rows.append([
                i.get("source_name", ""), i.get("type", ""),
                "Active" if active else "Inactive",
                _fmt_money(net), str(months if active else 0), _fmt_money(total),
            ])
        rows.append(["", "", "", "", "Total:", _fmt_money(grand)])
        t = Table(rows, colWidths=[1.6 * inch, 0.9 * inch, 0.8 * inch, 1.1 * inch, 0.8 * inch, 1.1 * inch])
        ts = _table_style()
        ts.add("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold")
        ts.add("BACKGROUND", (0, -1), (-1, -1), BRAND_LIGHT)
        t.setStyle(ts)
        story.append(t)

    elif report_type == "statement_expenses":
        _pdf_header(story, user, start, end, "Expense Statement", "All recurring expenses for the selected period.")
        story.append(Paragraph("Monthly Expenses", section))
        rows = [["Vendor", "Category", "Monthly", "× Months", "Period Total", "Auto-Pay"]]
        grand = 0.0
        for e in expenses:
            amt = float(e.get("monthly_amount") or 0)
            total = amt * months
            grand += total
            rows.append([
                e.get("vendor", ""), e.get("category", ""),
                _fmt_money(amt), str(months), _fmt_money(total),
                "Yes" if e.get("auto_pay") else "No",
            ])
        rows.append(["", "", "", "Total:", _fmt_money(grand), ""])
        t = Table(rows, colWidths=[1.6 * inch, 1.2 * inch, 1.0 * inch, 0.7 * inch, 1.2 * inch, 0.7 * inch])
        ts = _table_style()
        ts.add("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold")
        ts.add("BACKGROUND", (0, -1), (-1, -1), BRAND_LIGHT)
        t.setStyle(ts)
        story.append(t)

    elif report_type == "snapshot":
        _pdf_header(story, user, start, end, "Financial Snapshot", "A polished one-page overview of your finances.")
        # Top KPI grid
        kpi = [
            ["Total Income", _fmt_money(agg["period_income"])],
            ["Total Expenses", _fmt_money(agg["period_expense"])],
            ["Net Surplus", _fmt_money(agg["period_surplus"])],
            ["Net Worth", _fmt_money(agg["net_worth"])],
        ]
        kpi_table = Table([[k[0] for k in kpi], [k[1] for k in kpi]], colWidths=[1.6 * inch] * 4)
        kpi_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), BRAND_BLUE),
            ("TEXTCOLOR", (0, 0), (-1, 0), rl_colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 9),
            ("FONTNAME", (0, 1), (-1, 1), "Helvetica-Bold"),
            ("FONTSIZE", (0, 1), (-1, 1), 14),
            ("TEXTCOLOR", (0, 1), (-1, 1), BRAND_DARK),
            ("BACKGROUND", (0, 1), (-1, 1), BRAND_LIGHT),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("TOPPADDING", (0, 0), (-1, -1), 8),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ]))
        story.append(kpi_table)

        story.append(Paragraph("Income vs Expenses (Period)", section))
        story.append(_bar_chart_table(agg["period_income"], agg["period_expense"]))

        story.append(Paragraph("Expenses by Category (Monthly)", section))
        if agg["by_category"]:
            cat_rows = [["Category", "Monthly", "% of Outflow"]]
            total = agg["monthly_expense"] or 1
            for cat, val in sorted(agg["by_category"].items(), key=lambda x: -x[1]):
                cat_rows.append([cat, _fmt_money(val), _fmt_pct((val / total) * 100)])
            t = Table(cat_rows, colWidths=[2.5 * inch, 1.5 * inch, 1.5 * inch])
            t.setStyle(_table_style())
            story.append(t)
        else:
            story.append(Paragraph("No expenses recorded.", body))

        story.append(Paragraph("Debt Summary", section))
        story.append(Paragraph(
            f"Total debt: <b>{_fmt_money(agg['total_debt'])}</b> across {len(debts)} account(s).",
            body))
        story.append(Paragraph("Net Worth Summary", section))
        story.append(Paragraph(
            f"Assets: <b>{_fmt_money(agg['total_assets'])}</b> &nbsp;–&nbsp; Debts: <b>{_fmt_money(agg['total_debt'])}</b> &nbsp;=&nbsp; Net Worth: <b>{_fmt_money(agg['net_worth'])}</b>",
            body))

    else:  # detailed
        _pdf_header(story, user, start, end, "Detailed Financial Report",
                    "A complete picture of your finances — for advisors, banks, and lenders.")
        # 1. Summary
        story.append(Paragraph("1. Period Summary", section))
        sm = [
            ["Total Income", _fmt_money(agg["period_income"])],
            ["Total Expenses", _fmt_money(agg["period_expense"])],
            ["Net Surplus", _fmt_money(agg["period_surplus"])],
            ["Total Debt", _fmt_money(agg["total_debt"])],
            ["Total Assets", _fmt_money(agg["total_assets"])],
            ["Net Worth", _fmt_money(agg["net_worth"])],
        ]
        t = Table(sm, colWidths=[3.0 * inch, 2.0 * inch])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (0, -1), BRAND_LIGHT),
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("GRID", (0, 0), (-1, -1), 0.25, BRAND_GREY),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("ALIGN", (1, 0), (1, -1), "RIGHT"),
        ]))
        story.append(t)

        # 2. Income
        story.append(Paragraph("2. All Income Sources", section))
        if income:
            rows = [["Source", "Type", "Status", "Monthly Net", "Period Total"]]
            for i in income:
                net = float(i.get("net_monthly") or 0)
                active = i.get("is_active", True)
                rows.append([i.get("source_name", ""), i.get("type", ""),
                             "Active" if active else "Inactive",
                             _fmt_money(net), _fmt_money(net * months if active else 0)])
            t = Table(rows, colWidths=[2.0 * inch, 1.0 * inch, 0.9 * inch, 1.2 * inch, 1.2 * inch])
            t.setStyle(_table_style())
            story.append(t)
        else:
            story.append(Paragraph("No income sources on file.", body))

        # 3. Expenses
        story.append(Paragraph("3. All Expenses (Itemized)", section))
        if expenses:
            rows = [["Vendor", "Category", "Monthly", "Period Total", "Due Day", "Auto-Pay"]]
            for e in expenses:
                amt = float(e.get("monthly_amount") or 0)
                rows.append([e.get("vendor", ""), e.get("category", ""),
                             _fmt_money(amt), _fmt_money(amt * months),
                             str(e.get("due_day_of_month", "")),
                             "Yes" if e.get("auto_pay") else "No"])
            t = Table(rows, colWidths=[1.6 * inch, 1.1 * inch, 1.0 * inch, 1.1 * inch, 0.7 * inch, 0.8 * inch])
            t.setStyle(_table_style())
            story.append(t)

        # 4. Debt
        story.append(PageBreak())
        story.append(Paragraph("4. Debt Accounts", section))
        if debts:
            rows = [["Lender", "Type", "Balance", "APR", "Min Payment"]]
            for d in debts:
                rows.append([d.get("lender", ""), d.get("debt_type", ""),
                             _fmt_money(float(d.get("balance") or 0)),
                             _fmt_pct(float(d.get("apr") or 0) * 100 if (d.get("apr") or 0) < 1 else float(d.get("apr") or 0)),
                             _fmt_money(float(d.get("minimum_payment") or 0))])
            t = Table(rows, colWidths=[1.8 * inch, 1.2 * inch, 1.3 * inch, 0.9 * inch, 1.3 * inch])
            t.setStyle(_table_style())
            story.append(t)
        else:
            story.append(Paragraph("No debt accounts on file.", body))

        # 5. Assets
        story.append(Paragraph("5. Assets", section))
        if assets:
            rows = [["Name", "Type", "Current Value", "Purchase Value"]]
            for a in assets:
                rows.append([a.get("name", ""), a.get("asset_type", ""),
                             _fmt_money(float(a.get("current_value") or 0)),
                             _fmt_money(float(a.get("purchase_value") or 0))])
            t = Table(rows, colWidths=[2.0 * inch, 1.5 * inch, 1.5 * inch, 1.5 * inch])
            t.setStyle(_table_style())
            story.append(t)

        # 6. Investments
        story.append(Paragraph("6. Investments & Retirement", section))
        if investments:
            rows = [["Account", "Type", "Balance", "Monthly Contrib.", "Projected @ 65"]]
            for i in investments:
                rows.append([i.get("nickname") or i.get("type", ""), i.get("type", ""),
                             _fmt_money(float(i.get("balance") or 0)),
                             _fmt_money(float(i.get("contribution_monthly") or 0)),
                             _fmt_money(float(i.get("projected_at_65") or 0))])
            t = Table(rows, colWidths=[1.6 * inch, 1.2 * inch, 1.2 * inch, 1.4 * inch, 1.3 * inch])
            t.setStyle(_table_style())
            story.append(t)

        # 7. Financial Health
        score = int(user.get("financial_health_score") or 0)
        story.append(Paragraph("7. Financial Health Score", section))
        story.append(Paragraph(
            f"<b>Score: {score}/100</b><br/>"
            f"Calculated from your debt-to-income ratio, emergency fund coverage, savings rate, "
            f"investment allocation, and net worth growth trajectory. A higher score indicates "
            f"better long-term financial resilience.",
            body))

    doc.build(story, onFirstPage=_pdf_footer, onLaterPages=_pdf_footer)
    return buf.getvalue()


# ============================ DOCX ========================================
def _docx_heading(doc, text, size=22, color=DOCX_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = color
    return p


def _docx_brand_header(doc, user, start, end, title_text):
    p = doc.add_paragraph()
    r = p.add_run("PLOS — Personal Life Operating System")
    r.font.size = Pt(9)
    r.font.color.rgb = DOCX_GREY

    p = doc.add_paragraph()
    r = p.add_run(title_text)
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = DOCX_DARK
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    r = p.add_run(f"Prepared for: {user.get('full_name', 'User')}  |  {user.get('email', '')}")
    r.font.size = Pt(10)
    r.font.color.rgb = DOCX_GREY

    p = doc.add_paragraph()
    r = p.add_run(f"Period: {_date_range_label(start, end)}   ·   Generated {datetime.utcnow().strftime('%b %d, %Y %H:%M UTC')}")
    r.font.size = Pt(9)
    r.font.color.rgb = DOCX_GREY
    p.paragraph_format.space_after = Pt(12)


def _docx_table(doc, headers, rows, totals_row=None):
    table = doc.add_table(rows=1 + len(rows) + (1 if totals_row else 0), cols=len(headers))
    table.style = "Light Grid Accent 1"

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # apply blue shading via xml
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for cell in hdr:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1E40AF")
        tc_pr.append(shd)

    for ri, row in enumerate(rows):
        cells = table.rows[1 + ri].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    if totals_row:
        cells = table.rows[-1].cells
        for ci, val in enumerate(totals_row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.bold = True
    return table


def _gen_docx(report_type, start, end, months, user, income, expenses, debts, assets, investments) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.8)
    agg = _aggregate(income, expenses, debts, assets, months)

    if report_type == "statement_income":
        _docx_brand_header(doc, user, start, end, "Income Statement")
        _docx_heading(doc, "Income Sources", size=13)
        rows = []
        grand = 0.0
        for i in income:
            net = float(i.get("net_monthly") or 0)
            active = i.get("is_active", True)
            total = net * months if active else 0
            grand += total
            rows.append([i.get("source_name", ""), i.get("type", ""),
                         "Active" if active else "Inactive",
                         _fmt_money(net), str(months if active else 0), _fmt_money(total)])
        _docx_table(doc, ["Source", "Type", "Status", "Monthly Net", "× Months", "Period Total"],
                    rows, ["", "", "", "", "Total:", _fmt_money(grand)])

    elif report_type == "statement_expenses":
        _docx_brand_header(doc, user, start, end, "Expense Statement")
        _docx_heading(doc, "Monthly Expenses", size=13)
        rows = []
        grand = 0.0
        for e in expenses:
            amt = float(e.get("monthly_amount") or 0)
            total = amt * months
            grand += total
            rows.append([e.get("vendor", ""), e.get("category", ""),
                         _fmt_money(amt), str(months), _fmt_money(total),
                         "Yes" if e.get("auto_pay") else "No"])
        _docx_table(doc, ["Vendor", "Category", "Monthly", "× Months", "Period Total", "Auto-Pay"],
                    rows, ["", "", "", "Total:", _fmt_money(grand), ""])

    elif report_type == "snapshot":
        _docx_brand_header(doc, user, start, end, "Financial Snapshot")
        _docx_heading(doc, "Headline Numbers", size=13)
        _docx_table(doc, ["Metric", "Value"], [
            ["Total Income", _fmt_money(agg["period_income"])],
            ["Total Expenses", _fmt_money(agg["period_expense"])],
            ["Net Surplus", _fmt_money(agg["period_surplus"])],
            ["Total Debt", _fmt_money(agg["total_debt"])],
            ["Total Assets", _fmt_money(agg["total_assets"])],
            ["Net Worth", _fmt_money(agg["net_worth"])],
        ])
        _docx_heading(doc, "Expenses by Category (Monthly)", size=13)
        total = agg["monthly_expense"] or 1
        cat_rows = [[cat, _fmt_money(val), _fmt_pct((val / total) * 100)]
                    for cat, val in sorted(agg["by_category"].items(), key=lambda x: -x[1])]
        if cat_rows:
            _docx_table(doc, ["Category", "Monthly", "% of Outflow"], cat_rows)

    else:  # detailed
        _docx_brand_header(doc, user, start, end, "Detailed Financial Report")
        _docx_heading(doc, "1. Period Summary", size=13)
        _docx_table(doc, ["Metric", "Value"], [
            ["Total Income", _fmt_money(agg["period_income"])],
            ["Total Expenses", _fmt_money(agg["period_expense"])],
            ["Net Surplus", _fmt_money(agg["period_surplus"])],
            ["Total Debt", _fmt_money(agg["total_debt"])],
            ["Total Assets", _fmt_money(agg["total_assets"])],
            ["Net Worth", _fmt_money(agg["net_worth"])],
        ])
        _docx_heading(doc, "2. All Income Sources", size=13)
        if income:
            _docx_table(doc, ["Source", "Type", "Status", "Monthly Net", "Period Total"],
                        [[i.get("source_name", ""), i.get("type", ""),
                          "Active" if i.get("is_active", True) else "Inactive",
                          _fmt_money(float(i.get("net_monthly") or 0)),
                          _fmt_money(float(i.get("net_monthly") or 0) * months if i.get("is_active", True) else 0)]
                         for i in income])
        _docx_heading(doc, "3. All Expenses (Itemized)", size=13)
        if expenses:
            _docx_table(doc, ["Vendor", "Category", "Monthly", "Period Total", "Due Day", "Auto-Pay"],
                        [[e.get("vendor", ""), e.get("category", ""),
                          _fmt_money(float(e.get("monthly_amount") or 0)),
                          _fmt_money(float(e.get("monthly_amount") or 0) * months),
                          str(e.get("due_day_of_month", "")),
                          "Yes" if e.get("auto_pay") else "No"] for e in expenses])
        _docx_heading(doc, "4. Debt Accounts", size=13)
        if debts:
            _docx_table(doc, ["Lender", "Type", "Balance", "APR", "Min Payment"],
                        [[d.get("lender", ""), d.get("debt_type", ""),
                          _fmt_money(float(d.get("balance") or 0)),
                          _fmt_pct(float(d.get("apr") or 0) * 100 if (d.get("apr") or 0) < 1 else float(d.get("apr") or 0)),
                          _fmt_money(float(d.get("minimum_payment") or 0))] for d in debts])
        _docx_heading(doc, "5. Assets", size=13)
        if assets:
            _docx_table(doc, ["Name", "Type", "Current Value", "Purchase Value"],
                        [[a.get("name", ""), a.get("asset_type", ""),
                          _fmt_money(float(a.get("current_value") or 0)),
                          _fmt_money(float(a.get("purchase_value") or 0))] for a in assets])
        _docx_heading(doc, "6. Investments & Retirement", size=13)
        if investments:
            _docx_table(doc, ["Account", "Type", "Balance", "Monthly Contrib.", "Projected @ 65"],
                        [[i.get("nickname") or i.get("type", ""), i.get("type", ""),
                          _fmt_money(float(i.get("balance") or 0)),
                          _fmt_money(float(i.get("contribution_monthly") or 0)),
                          _fmt_money(float(i.get("projected_at_65") or 0))] for i in investments])
        _docx_heading(doc, "7. Financial Health Score", size=13)
        score = int(user.get("financial_health_score") or 0)
        p = doc.add_paragraph()
        r = p.add_run(f"Score: {score}/100\n")
        r.font.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = DOCX_DARK
        r2 = p.add_run("Calculated from your debt-to-income ratio, emergency fund coverage, savings "
                       "rate, investment allocation, and net worth growth trajectory.")
        r2.font.size = Pt(10)
        r2.font.color.rgb = DOCX_GREY

    # Footer with PLOS branding
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("PLOS Financial Report  ·  Confidential")
    r.font.size = Pt(8)
    r.font.color.rgb = DOCX_GREY

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
