"""PLOS Career — Resume Vault module.

Stores multiple user-uploaded resumes (PDF, DOCX, DOC, TXT, or pasted text).
The user picks one as the "base" for AI-tailored resume + cover letter
generation via career_tailor.py.

Collection: user_resumes
{
  resume_id, user_id, name, file_type ("pdf"|"docx"|"txt"|"paste"),
  content_b64 (original file bytes, for future re-download),
  text (extracted plain text used for LLM prompting),
  is_default (bool), size_bytes, uploaded_at
}
"""
from __future__ import annotations

import base64
import io
import logging
import uuid
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, Depends, HTTPException
from pydantic import BaseModel, Field

logger = logging.getLogger(__name__)

ALLOWED_EXT = {"pdf", "docx", "doc", "txt", "paste"}
MAX_BYTES = 5 * 1024 * 1024  # 5 MB per resume


# ---------- Text extraction ------------------------------------------------
def _extract_pdf_text(raw: bytes) -> str:
    try:
        import pdfplumber
        out = []
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                if t.strip():
                    out.append(t)
        return "\n\n".join(out).strip()
    except Exception as exc:
        logger.warning("PDF extraction failed: %s", exc)
        return ""


def _extract_docx_text(raw: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(raw))
        parts: List[str] = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return "\n".join(parts).strip()
    except Exception as exc:
        logger.warning("DOCX extraction failed: %s", exc)
        return ""


def extract_resume_text(raw_b64: str, file_type: str) -> str:
    """Decode + extract plain text from a base64 file blob."""
    try:
        raw = base64.b64decode(raw_b64)
    except Exception:
        return ""
    ft = (file_type or "").lower()
    if ft == "pdf":
        return _extract_pdf_text(raw)
    if ft in ("docx", "doc"):
        # python-docx also opens most .doc files (best-effort)
        return _extract_docx_text(raw)
    if ft == "txt":
        try:
            return raw.decode("utf-8", errors="ignore").strip()
        except Exception:
            return ""
    return ""


# ---------- Models ---------------------------------------------------------
class ResumeCreate(BaseModel):
    name: str = Field(..., min_length=1, max_length=120)
    file_type: str = Field(..., description="pdf | docx | doc | txt | paste")
    content_b64: Optional[str] = None  # required for pdf/docx/doc/txt
    text: Optional[str] = None         # required for "paste"; auto-populated for others


class ResumeUpdate(BaseModel):
    name: Optional[str] = None
    text: Optional[str] = None
    is_default: Optional[bool] = None


# ---------- Router factory -------------------------------------------------
def make_router(db, get_current_user_id):
    r = APIRouter(prefix="/api/career/resumes", tags=["career-resumes"])

    @r.get("")
    async def list_resumes(user_id: str = Depends(get_current_user_id)):
        docs = await db.user_resumes.find(
            {"user_id": user_id}, {"_id": 0, "content_b64": 0}
        ).sort("uploaded_at", -1).to_list(50)
        return {"resumes": docs}

    @r.get("/{resume_id}")
    async def get_resume(
        resume_id: str, user_id: str = Depends(get_current_user_id)
    ):
        doc = await db.user_resumes.find_one(
            {"user_id": user_id, "resume_id": resume_id}, {"_id": 0}
        )
        if not doc:
            raise HTTPException(404, "Resume not found")
        return doc

    @r.post("", status_code=201)
    async def create_resume(
        body: ResumeCreate, user_id: str = Depends(get_current_user_id)
    ):
        ft = body.file_type.lower()
        if ft not in ALLOWED_EXT:
            raise HTTPException(400, f"file_type must be one of {sorted(ALLOWED_EXT)}")

        content_b64 = body.content_b64 or ""
        text = (body.text or "").strip()

        if ft == "paste":
            if not text:
                raise HTTPException(400, "text is required for paste")
        else:
            if not content_b64:
                raise HTTPException(400, "content_b64 is required for file upload")
            # size check
            try:
                raw_len = len(base64.b64decode(content_b64, validate=False))
            except Exception:
                raise HTTPException(400, "content_b64 is not valid base64")
            if raw_len > MAX_BYTES:
                raise HTTPException(
                    413, f"File too large (max {MAX_BYTES // (1024 * 1024)} MB)"
                )
            if not text:
                text = extract_resume_text(content_b64, ft)
            if not text:
                # Best-effort — still save the file but flag empty extraction.
                text = ""

        resume_id = f"res_{uuid.uuid4().hex[:12]}"
        now = datetime.now(timezone.utc).isoformat()
        doc: Dict[str, Any] = {
            "resume_id": resume_id,
            "user_id": user_id,
            "name": body.name.strip(),
            "file_type": ft,
            "content_b64": content_b64 if ft != "paste" else "",
            "text": text,
            "size_bytes": len(text.encode("utf-8")),
            "is_default": False,
            "uploaded_at": now,
        }
        # Auto-set as default if this is the user's first resume
        existing_count = await db.user_resumes.count_documents({"user_id": user_id})
        if existing_count == 0:
            doc["is_default"] = True

        await db.user_resumes.insert_one(doc)
        # Return without the heavy blob
        out = {**doc}
        out.pop("_id", None)
        out.pop("content_b64", None)
        return out

    @r.put("/{resume_id}")
    async def update_resume(
        resume_id: str,
        body: ResumeUpdate,
        user_id: str = Depends(get_current_user_id),
    ):
        update: Dict[str, Any] = {}
        if body.name is not None:
            update["name"] = body.name.strip()
        if body.text is not None:
            update["text"] = body.text.strip()
            update["size_bytes"] = len(update["text"].encode("utf-8"))
        if body.is_default is True:
            # Clear existing defaults first
            await db.user_resumes.update_many(
                {"user_id": user_id, "is_default": True},
                {"$set": {"is_default": False}},
            )
            update["is_default"] = True
        if not update:
            raise HTTPException(400, "No editable fields provided")

        res = await db.user_resumes.update_one(
            {"user_id": user_id, "resume_id": resume_id}, {"$set": update}
        )
        if res.matched_count == 0:
            raise HTTPException(404, "Resume not found")
        return {"ok": True, "updated": list(update.keys())}

    @r.delete("/{resume_id}")
    async def delete_resume(
        resume_id: str, user_id: str = Depends(get_current_user_id)
    ):
        doc = await db.user_resumes.find_one(
            {"user_id": user_id, "resume_id": resume_id}
        )
        if not doc:
            raise HTTPException(404, "Resume not found")
        await db.user_resumes.delete_one({"user_id": user_id, "resume_id": resume_id})
        # If we removed the default, promote the newest remaining resume.
        if doc.get("is_default"):
            newest = await db.user_resumes.find_one(
                {"user_id": user_id}, sort=[("uploaded_at", -1)]
            )
            if newest:
                await db.user_resumes.update_one(
                    {"_id": newest["_id"]}, {"$set": {"is_default": True}}
                )
        return {"ok": True}

    return r


# ---------- Public helper (used by career_tailor.py) ---------------------
async def get_default_or_by_id(
    db, user_id: str, resume_id: Optional[str] = None
) -> Optional[Dict[str, Any]]:
    """Return the resume the tailor should use.

    If resume_id is provided, return that. Otherwise return the user's
    is_default=True resume, or the newest one as a last-resort fallback.
    """
    if resume_id:
        return await db.user_resumes.find_one(
            {"user_id": user_id, "resume_id": resume_id}, {"_id": 0}
        )
    doc = await db.user_resumes.find_one(
        {"user_id": user_id, "is_default": True}, {"_id": 0}
    )
    if doc:
        return doc
    # Last resort — newest
    return await db.user_resumes.find_one(
        {"user_id": user_id}, {"_id": 0}, sort=[("uploaded_at", -1)]
    )
