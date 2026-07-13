"""
PLOS — Career file management
- Upload resume / job description / other career documents (PDF, DOCX, TXT)
- Extract text from uploads for AI analysis
- List & delete uploaded files
- Create resume from structured form (Claude-generated polished text)
- Download a generated resume as PDF or DOCX
"""
from __future__ import annotations

import base64
import io
import uuid
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from pydantic import BaseModel, Field

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from reportlab.lib.colors import HexColor

router = APIRouter(prefix="/api/career-files", tags=["career-files"])

MAX_FILE_SIZE = 5 * 1024 * 1024  # 5 MB

ALLOWED_KINDS = {"resume", "job_description", "other"}
ALLOWED_MIME = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
    "text/plain",
    "image/jpeg",  # accepted for "other"
    "image/png",
}


# --------------------------- Models ---------------------------------------
class ResumeWorkEntry(BaseModel):
    title: str
    employer: str
    location: Optional[str] = None
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    bullets: List[str] = []


class ResumeEducation(BaseModel):
    degree: str
    institution: str
    year: Optional[str] = None
    notes: Optional[str] = None


class ResumeCertification(BaseModel):
    name: str
    issuer: Optional[str] = None
    year: Optional[str] = None


class ResumeDraft(BaseModel):
    full_name: str
    email: Optional[str] = None
    phone: Optional[str] = None
    location: Optional[str] = None
    summary: Optional[str] = None
    work_experience: List[ResumeWorkEntry] = []
    education: List[ResumeEducation] = []
    skills: List[str] = []
    certifications: List[ResumeCertification] = []
    awards: List[str] = []


class ResumeDraftSave(BaseModel):
    draft: ResumeDraft


# --------------------------- Helpers --------------------------------------
def _strip(doc: Optional[dict]) -> Optional[dict]:
    if not doc:
        return None
    doc.pop("_id", None)
    return doc


def _extract_text(content: bytes, mime: str) -> str:
    """Best-effort text extraction. Returns empty string if not extractable."""
    try:
        if mime == "application/pdf":
            reader = PdfReader(io.BytesIO(content))
            return "\n".join((p.extract_text() or "") for p in reader.pages).strip()
        if mime in {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        }:
            doc = Document(io.BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs).strip()
        if mime == "text/plain":
            return content.decode("utf-8", errors="ignore").strip()
    except Exception:
        return ""
    return ""


# --------------------------- PDF / DOCX generation -------------------------
BRAND_BLUE = HexColor("#1E40AF")
BRAND_DARK = HexColor("#0F172A")
BRAND_GREY = HexColor("#64748B")
DOCX_BLUE = RGBColor(0x1E, 0x40, 0xAF)
DOCX_GREY = RGBColor(0x64, 0x75, 0x8B)
DOCX_DARK = RGBColor(0x0F, 0x17, 0x23)


def _resume_to_pdf(draft: Dict[str, Any]) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER,
                            leftMargin=0.6 * inch, rightMargin=0.6 * inch,
                            topMargin=0.6 * inch, bottomMargin=0.7 * inch,
                            title=f"Resume — {draft.get('full_name', 'User')}")
    base = getSampleStyleSheet()
    name_style = ParagraphStyle("Name", parent=base["Title"], fontSize=22, textColor=BRAND_DARK,
                                fontName="Helvetica-Bold", leading=26, spaceAfter=4)
    contact = ParagraphStyle("Contact", parent=base["Normal"], fontSize=9,
                             textColor=BRAND_GREY, spaceAfter=12)
    section = ParagraphStyle("Section", parent=base["Heading2"], fontSize=12, textColor=BRAND_BLUE,
                             spaceBefore=10, spaceAfter=4, fontName="Helvetica-Bold")
    body = ParagraphStyle("Body", parent=base["Normal"], fontSize=10, textColor=BRAND_DARK, leading=13)
    job = ParagraphStyle("Job", parent=body, fontName="Helvetica-Bold", fontSize=10, spaceBefore=4)
    job_meta = ParagraphStyle("JobMeta", parent=base["Normal"], fontSize=9, textColor=BRAND_GREY)
    bullet = ParagraphStyle("Bullet", parent=body, leftIndent=12, bulletIndent=4)

    story = []
    story.append(Paragraph(draft.get("full_name", "Your Name"), name_style))
    contact_parts = [draft.get("email"), draft.get("phone"), draft.get("location")]
    contact_line = "  •  ".join(p for p in contact_parts if p)
    if contact_line:
        story.append(Paragraph(contact_line, contact))

    if draft.get("summary"):
        story.append(Paragraph("Summary", section))
        story.append(Paragraph(draft["summary"], body))

    if draft.get("work_experience"):
        story.append(Paragraph("Experience", section))
        for w in draft["work_experience"]:
            story.append(Paragraph(f"{w.get('title','')} — {w.get('employer','')}", job))
            meta = []
            if w.get("location"):
                meta.append(w["location"])
            date = " – ".join(d for d in [w.get("start_date"), w.get("end_date")] if d)
            if date:
                meta.append(date)
            if meta:
                story.append(Paragraph("  •  ".join(meta), job_meta))
            for b in (w.get("bullets") or []):
                story.append(Paragraph("• " + b, bullet))

    if draft.get("education"):
        story.append(Paragraph("Education", section))
        for e in draft["education"]:
            line = f"<b>{e.get('degree','')}</b> — {e.get('institution','')}"
            if e.get("year"):
                line += f"  ({e['year']})"
            story.append(Paragraph(line, body))
            if e.get("notes"):
                story.append(Paragraph(e["notes"], body))

    if draft.get("skills"):
        story.append(Paragraph("Skills", section))
        story.append(Paragraph(" • ".join(draft["skills"]), body))

    if draft.get("certifications"):
        story.append(Paragraph("Certifications", section))
        for c in draft["certifications"]:
            line = f"<b>{c.get('name','')}</b>"
            if c.get("issuer"):
                line += f" — {c['issuer']}"
            if c.get("year"):
                line += f" ({c['year']})"
            story.append(Paragraph(line, body))

    if draft.get("awards"):
        story.append(Paragraph("Awards", section))
        for a in draft["awards"]:
            story.append(Paragraph("• " + a, bullet))

    doc.build(story)
    return buf.getvalue()


def _resume_to_docx(draft: Dict[str, Any]) -> bytes:
    d = Document()
    sec = d.sections[0]
    sec.left_margin = Inches(0.7)
    sec.right_margin = Inches(0.7)
    sec.top_margin = Inches(0.6)
    sec.bottom_margin = Inches(0.7)

    def heading(text, size=12, color=DOCX_BLUE):
        p = d.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = True
        r.font.color.rgb = color

    # Name
    p = d.add_paragraph()
    r = p.add_run(draft.get("full_name", "Your Name"))
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = DOCX_DARK

    contact_parts = [draft.get("email"), draft.get("phone"), draft.get("location")]
    contact = "  •  ".join(p for p in contact_parts if p)
    if contact:
        p = d.add_paragraph()
        r = p.add_run(contact)
        r.font.size = Pt(9)
        r.font.color.rgb = DOCX_GREY

    if draft.get("summary"):
        heading("Summary")
        d.add_paragraph(draft["summary"])

    if draft.get("work_experience"):
        heading("Experience")
        for w in draft["work_experience"]:
            p = d.add_paragraph()
            r = p.add_run(f"{w.get('title','')} — {w.get('employer','')}")
            r.font.bold = True
            r.font.size = Pt(11)
            meta = " · ".join(m for m in [w.get("location"), " – ".join(x for x in [w.get("start_date"), w.get("end_date")] if x)] if m)
            if meta:
                p2 = d.add_paragraph()
                r2 = p2.add_run(meta)
                r2.font.size = Pt(9)
                r2.font.color.rgb = DOCX_GREY
            for b in (w.get("bullets") or []):
                d.add_paragraph("• " + b)

    if draft.get("education"):
        heading("Education")
        for e in draft["education"]:
            line = f"{e.get('degree','')} — {e.get('institution','')}"
            if e.get("year"):
                line += f" ({e['year']})"
            d.add_paragraph(line)
            if e.get("notes"):
                d.add_paragraph(e["notes"])

    if draft.get("skills"):
        heading("Skills")
        d.add_paragraph(" • ".join(draft["skills"]))

    if draft.get("certifications"):
        heading("Certifications")
        for c in draft["certifications"]:
            parts = [c.get("name", "")]
            if c.get("issuer"):
                parts.append(c["issuer"])
            if c.get("year"):
                parts.append(c["year"])
            d.add_paragraph(" — ".join(parts))

    if draft.get("awards"):
        heading("Awards")
        for a in draft["awards"]:
            d.add_paragraph("• " + a)

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


# --------------------------- Factory --------------------------------------
def make_router(db, get_current_user_id, emergent_llm_key, llm_chat_cls, user_msg_cls):
    async def call_claude(session_id: str, system: str, prompt: str) -> str:
        chat = llm_chat_cls(
            api_key=emergent_llm_key, session_id=session_id, system_message=system,
        ).with_model("anthropic", "claude-sonnet-4-5-20250929")
        resp = await chat.send_message(user_msg_cls(text=prompt))
        return resp if isinstance(resp, str) else str(resp)

    # ----------- List & Delete uploaded files -----------
    @router.get("/list")
    async def list_files(
        kind: Optional[str] = None, user_id: str = Depends(get_current_user_id)
    ):
        q: Dict[str, Any] = {"user_id": user_id}
        if kind:
            if kind not in ALLOWED_KINDS:
                raise HTTPException(status_code=400, detail=f"Invalid kind: {kind}")
            q["kind"] = kind
        items = []
        async for f in db.career_files.find(q).sort("uploaded_at", -1):
            f = _strip(f)
            # don't return content / text_content in list view
            f.pop("content_b64", None)
            f.pop("text_content", None)
            items.append(f)
        return {"files": items}

    @router.get("/file/{file_id}")
    async def get_file_meta(file_id: str, user_id: str = Depends(get_current_user_id)):
        f = await db.career_files.find_one({"user_id": user_id, "file_id": file_id})
        if not f:
            raise HTTPException(status_code=404, detail="File not found")
        f = _strip(f)
        # return text_content but not the raw bytes for inline preview
        f.pop("content_b64", None)
        return f

    @router.get("/file/{file_id}/download")
    async def download_file(file_id: str, user_id: str = Depends(get_current_user_id)):
        f = await db.career_files.find_one({"user_id": user_id, "file_id": file_id})
        if not f:
            raise HTTPException(status_code=404, detail="File not found")
        return {
            "filename": f.get("filename", "file"),
            "mime_type": f.get("mime", "application/octet-stream"),
            "content_base64": f.get("content_b64", ""),
            "size_bytes": f.get("size", 0),
        }

    @router.delete("/file/{file_id}")
    async def delete_file(file_id: str, user_id: str = Depends(get_current_user_id)):
        r = await db.career_files.delete_one({"user_id": user_id, "file_id": file_id})
        if r.deleted_count == 0:
            raise HTTPException(status_code=404, detail="File not found")
        return {"ok": True}

    # ----------- Upload (multipart) -----------
    @router.post("/upload")
    async def upload_file(
        file: UploadFile = File(...),
        kind: str = Form(...),
        label: Optional[str] = Form(None),
        user_id: str = Depends(get_current_user_id),
    ):
        if kind not in ALLOWED_KINDS:
            raise HTTPException(status_code=400, detail=f"Invalid kind: {kind}")
        content = await file.read()
        if len(content) == 0:
            raise HTTPException(status_code=400, detail="Empty file")
        if len(content) > MAX_FILE_SIZE:
            raise HTTPException(status_code=400, detail=f"File too large (max {MAX_FILE_SIZE // 1024 // 1024} MB)")

        mime = file.content_type or "application/octet-stream"
        # Some browsers send empty/wrong content_type — fall back to filename ext
        if mime == "application/octet-stream":
            fn = (file.filename or "").lower()
            if fn.endswith(".pdf"):
                mime = "application/pdf"
            elif fn.endswith(".docx"):
                mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            elif fn.endswith(".txt"):
                mime = "text/plain"

        # extract text best-effort
        text = _extract_text(content, mime)
        file_id = str(uuid.uuid4())
        rec = {
            "file_id": file_id,
            "user_id": user_id,
            "kind": kind,
            "filename": file.filename or "untitled",
            "mime": mime,
            "size": len(content),
            "label": label,
            "uploaded_at": datetime.now(timezone.utc).isoformat(),
            "text_content": text[:200_000],  # cap stored text to avoid bloat
            "content_b64": base64.b64encode(content).decode("ascii"),
        }
        await db.career_files.insert_one(rec)
        # surface short preview only
        rec.pop("content_b64", None)
        rec.pop("_id", None)
        rec["text_preview"] = (text[:600] + "…") if len(text) > 600 else text
        rec.pop("text_content", None)
        return {"file": rec}

    # ----------- Resume draft (structured form) -----------
    @router.get("/resume-draft")
    async def get_resume_draft(user_id: str = Depends(get_current_user_id)):
        d = await db.career_resume_drafts.find_one({"user_id": user_id})
        if not d:
            return {"draft": None}
        return {"draft": _strip(d).get("draft")}

    @router.put("/resume-draft")
    async def save_resume_draft(payload: ResumeDraftSave, user_id: str = Depends(get_current_user_id)):
        await db.career_resume_drafts.update_one(
            {"user_id": user_id},
            {"$set": {"draft": payload.draft.model_dump(),
                      "updated_at": datetime.now(timezone.utc).isoformat()}},
            upsert=True,
        )
        return {"ok": True}

    # ----------- Polish via Claude -----------
    @router.post("/resume-draft/polish")
    async def polish_resume(user_id: str = Depends(get_current_user_id)):
        d = await db.career_resume_drafts.find_one({"user_id": user_id})
        if not d or not d.get("draft"):
            raise HTTPException(status_code=400, detail="No resume draft saved yet")
        draft = d["draft"]
        system = "You are an executive resume writer. Output ONLY valid JSON."
        prompt = (
            "Rewrite the bullets in each work_experience entry below to be punchy, "
            "achievement-focused, and quantified where possible. Also tighten the 'summary' "
            "into a strong 2-3 sentence professional summary. Return JSON with the same shape "
            "(only summary and work_experience keys are required). Keep all factual claims; "
            "do not invent metrics. The user's data:\n" + str(draft)
        )
        try:
            text = await call_claude(f"resume-polish-{user_id}", system, prompt)
        except Exception as e:
            raise HTTPException(status_code=502, detail=f"PLOS AI error: {e}")
        # naive JSON extraction
        import json
        import re
        m = re.search(r"\{.*\}", text, re.DOTALL)
        if not m:
            raise HTTPException(status_code=502, detail="Unable to parse PLOS AI response")
        try:
            data = json.loads(m.group(0))
        except Exception:
            raise HTTPException(status_code=502, detail="Invalid JSON from PLOS AI")

        # merge improvements back into draft
        new_draft = {**draft}
        if data.get("summary"):
            new_draft["summary"] = data["summary"]
        if data.get("work_experience"):
            new_draft["work_experience"] = data["work_experience"]
        await db.career_resume_drafts.update_one(
            {"user_id": user_id},
            {"$set": {"draft": new_draft, "updated_at": datetime.now(timezone.utc).isoformat()}},
            upsert=True,
        )
        return {"draft": new_draft}

    # ----------- Download generated resume -----------
    @router.post("/resume-draft/download")
    async def download_resume(payload: Dict[str, Any], user_id: str = Depends(get_current_user_id)):
        fmt = payload.get("format", "pdf").lower()
        if fmt not in {"pdf", "docx"}:
            raise HTTPException(status_code=400, detail="Format must be pdf or docx")
        d = await db.career_resume_drafts.find_one({"user_id": user_id})
        if not d or not d.get("draft"):
            raise HTTPException(status_code=400, detail="No resume draft saved yet")
        draft = d["draft"]
        full_name = (draft.get("full_name") or "Resume").replace(" ", "_")
        ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        if fmt == "pdf":
            data = _resume_to_pdf(draft)
            return {
                "filename": f"Resume_{full_name}_{ts}.pdf",
                "mime_type": "application/pdf",
                "content_base64": base64.b64encode(data).decode("ascii"),
                "size_bytes": len(data),
            }
        data = _resume_to_docx(draft)
        return {
            "filename": f"Resume_{full_name}_{ts}.docx",
            "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "content_base64": base64.b64encode(data).decode("ascii"),
            "size_bytes": len(data),
        }

    return router
